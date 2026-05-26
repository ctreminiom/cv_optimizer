"""Output generation tools: DOCX and Markdown report writing."""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from crewai.tools import tool

__all__ = ["generate_docx_output", "write_job_report"]


@tool("generate_docx_output")
def generate_docx_output(template_path: str, adapted_cv_json: str,
                         candidate_profile_json: str, output_path: str) -> str:
    """Generates the final adapted CV as DOCX, using the master CV as template.

    Args:
        template_path: Path to the master CV file (.docx preferred — keeps
            fonts/styles; .pdf masters trigger a clean-DOCX build).
        adapted_cv_json: JSON-stringified AdaptedCV with sections, summary,
            keywords_injected, etc.
        candidate_profile_json: JSON-stringified CandidateProfile with
            full_name, contact, work_experience, education, skills, etc.
        output_path: Absolute path where the new .docx will be written
            (e.g. "output/Acme_Senior-PM/adapted_cv.docx").
    """
    try:
        from docx import Document
    except ImportError:
        return json.dumps({"error": "python-docx not installed"})

    def _parse(value, key):
        if isinstance(value, dict):
            return value
        try:
            return json.loads(value)
        except (json.JSONDecodeError, TypeError):
            try:
                import ast
                return ast.literal_eval(value)
            except Exception as e:
                raise ValueError(f"could not parse {key}: {e}")

    try:
        adapted_cv = _parse(adapted_cv_json, "adapted_cv_json")
        profile = _parse(candidate_profile_json, "candidate_profile_json")
        if not output_path:
            return json.dumps({"error": "output_path is required"})
    except Exception as e:
        return json.dumps({"error": f"invalid input: {e}"})

    try:
        # Decide whether to reuse the master as a template.
        # DOCX masters preserve fonts, headings, and styling.
        # PDF masters cannot be reused — we build a clean DOCX instead.
        template_is_docx = (
            template_path
            and Path(template_path).exists()
            and Path(template_path).suffix.lower() == ".docx"
        )

        if template_is_docx:
            try:
                doc = Document(template_path)
                # Wipe existing content; keep style definitions intact.
                for p in list(doc.paragraphs):
                    p._element.getparent().remove(p._element)
            except Exception:
                # If the DOCX is corrupted or password-protected, fall back.
                doc = Document()
        else:
            doc = Document()

        doc.add_heading(profile.get("full_name", ""), level=0)
        contact = profile.get("contact", {})
        contact_line = " | ".join(v for v in [contact.get("email"), contact.get("phone"),
                                              contact.get("linkedin"), contact.get("location")] if v)
        if contact_line:
            doc.add_paragraph(contact_line)

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(adapted_cv.get("summary", ""))

        for section in adapted_cv.get("sections", []):
            doc.add_heading(section["name"], level=1)
            if section.get("summary_text"):
                doc.add_paragraph(section["summary_text"])
            for bullet in section.get("bullets", []):
                if bullet.get("rewritten"):
                    doc.add_paragraph(bullet["rewritten"], style="List Bullet")

        if profile.get("education"):
            doc.add_heading("Education", level=1)
            for edu in profile["education"]:
                line = f"{edu.get('degree','')}, {edu.get('institution','')}"
                if edu.get("end_date"):
                    line += f" — {edu['end_date']}"
                doc.add_paragraph(line)

        if profile.get("skills"):
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(profile["skills"]))

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return json.dumps({"status": "ok", "output_path": output_path})
    except Exception as e:
        return json.dumps({"error": f"failed to generate DOCX: {e}"})


# ─────────────────────────────────────────────────────────────────────────────
# TOOL — Markdown report writer
# ─────────────────────────────────────────────────────────────────────────────
def _wordlevel_diff(before: str, after: str) -> str:
    """Render a single-line word-level diff using strikethrough on removed
    runs and bold on added runs. Reads much faster than a two-paragraph
    before/after stack when bullets are long. §1.4."""
    from difflib import SequenceMatcher
    a = (before or "").split()
    b = (after or "").split()
    if not a and not b:
        return ""
    sm = SequenceMatcher(a=a, b=b, autojunk=False)
    out: list[str] = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            chunk = " ".join(a[i1:i2])
            if chunk:
                out.append(chunk)
        elif tag == "delete":
            chunk = " ".join(a[i1:i2])
            if chunk:
                out.append(f"~~{chunk}~~")
        elif tag == "insert":
            chunk = " ".join(b[j1:j2])
            if chunk:
                out.append(f"**{chunk}**")
        elif tag == "replace":
            del_chunk = " ".join(a[i1:i2])
            ins_chunk = " ".join(b[j1:j2])
            if del_chunk:
                out.append(f"~~{del_chunk}~~")
            if ins_chunk:
                out.append(f"**{ins_chunk}**")
    return " ".join(out)


@tool("write_job_report")
def write_job_report(report_json: str, output_path: str) -> str:
    """Writes the per-job markdown report combining all evaluations,
    candidate profile, prioritized changes, and cover letter.

    Args:
        report_json: A JSON string of the full JobReport (must include keys
            like job, candidate_profile, evaluations, gap_analysis, ats_report,
            consolidated_feedback, adapted_cv, authenticity_report,
            mirroring_report, verification_report, interview_questions,
            cover_letter — any subset is OK).
        output_path: Absolute filesystem path where the .md file will be
            written (e.g. "output/Acme_Senior-PM/report.md").
    """
    try:
        if isinstance(report_json, dict):
            report = report_json
        else:
            try:
                report = json.loads(report_json)
            except json.JSONDecodeError:
                import ast
                report = ast.literal_eval(report_json)
        if "report" in report and isinstance(report["report"], dict):
            # Allow the legacy nested form {"report": {...}, "output_path": ...}
            if not output_path and "output_path" in report:
                output_path = report["output_path"]
            report = report["report"]
        if not output_path:
            return json.dumps({"error": "output_path is required"})
    except Exception as e:
        return json.dumps({"error": f"invalid report_json: {e}"})

    # Defensive coercion — the agent's JSON might leave nested objects as
    # strings or wrap lists at unexpected nesting depths. Treat anything
    # that isn't a dict/list as empty rather than crashing.
    def _D(v):
        return v if isinstance(v, dict) else ({} if v is None else (v if False else {}))
    def _L(v):
        if isinstance(v, list): return v
        return []
    def _LD(v):
        return [x for x in _L(v) if isinstance(x, dict)]

    job = _D(report.get("job"))
    candidate = _D(report.get("candidate_profile")) or _D(report.get("candidate"))
    enrichment = _D(report.get("enrichment"))
    strategic = _D(enrichment.get("strategic"))
    md: list[str] = []

    # ──────────────────────────────────────────────────────────────────────
    # §1.2 — Canonical blocker registry. Reviewer red flags + critical gaps
    # are heavily redundant ("degree mismatch" appears 8+ times in the legacy
    # report, each reviewer paraphrased differently). We use fuzzy matching
    # so different phrasings of the same finding collapse onto one entry.
    # ──────────────────────────────────────────────────────────────────────
    _STOPWORDS = {
        "the", "a", "an", "is", "in", "on", "of", "to", "for", "with",
        "and", "or", "as", "by", "be", "are", "this", "that", "her",
        "his", "their", "from", "at", "than", "not", "no", "but",
        "candidate", "she", "he", "her", "his", "have", "has", "had",
        "may", "could", "would", "will", "which", "while", "el", "la",
        "los", "las", "de", "del", "y", "en", "un", "una", "que",
    }

    def _significant_tokens(s: str) -> set:
        import re as _re
        x = (s or "").lower()
        # Strip the leading reviewer label that _collect_red_flags adds.
        x = _re.sub(r"^_[^_]+_:\s*", "", x)
        x = _re.sub(r"[^a-z0-9áéíóúñü ]+", " ", x)
        toks = [t for t in x.split() if len(t) > 2 and t not in _STOPWORDS]
        return set(toks)

    def _findings_match(a: str, b: str) -> bool:
        """Two findings represent the same blocker if they share enough
        significant tokens. Tuned to collapse 'degree in Marketing/PR not
        Systems Engineering' phrasings while keeping distinct blockers
        separate."""
        ta = _significant_tokens(a)
        tb = _significant_tokens(b)
        if not ta or not tb:
            return False
        overlap = len(ta & tb)
        smaller = min(len(ta), len(tb))
        # Jaccard-ish: need either a strong overlap ratio OR a high absolute count.
        ratio = overlap / smaller if smaller else 0.0
        return overlap >= 4 and ratio >= 0.45

    def _build_canonical_blockers() -> list[dict]:
        blockers: list[dict] = []
        def _merge_or_add(text: str, kind: str, flagger: str) -> None:
            if not text or not text.strip():
                return
            text = text.strip()
            for b in blockers:
                if _findings_match(b["text"], text):
                    if flagger not in b["flagged_by"]:
                        b["flagged_by"].append(flagger)
                    # Prefer the longer / more specific phrasing as canonical.
                    if len(text) > len(b["text"]) and len(text) < len(b["text"]) * 2:
                        b["text"] = text
                    return
            blockers.append({"text": text, "kind": kind, "flagged_by": [flagger]})

        for g in (_D(report.get("gap_analysis")).get("critical_gaps") or []):
            if isinstance(g, str):
                _merge_or_add(g, "gap", "Gap analysis")
        for ev in _LD(report.get("evaluations")):
            reviewer = ev.get("agent_role", "Reviewer")
            for f in (ev.get("red_flags") or []):
                if isinstance(f, str):
                    _merge_or_add(f, "red_flag", reviewer)

        # Sort: more flaggers first, then gaps before red_flags.
        blockers.sort(key=lambda b: (-len(b["flagged_by"]),
                                        0 if b["kind"] == "gap" else 1))
        return blockers

    canonical_blockers = _build_canonical_blockers()

    # Reviewer roster — used for the comparative table header.
    reviewer_roles = []
    for ev in _LD(report.get("evaluations")):
        r = ev.get("agent_role") or "Reviewer"
        if r not in reviewer_roles:
            reviewer_roles.append(r)

    # ──────────────────────────────────────────────────────────────────────
    # Helper: pull the highest-priority red flags from every evaluator
    # ──────────────────────────────────────────────────────────────────────
    def _collect_red_flags() -> list[str]:
        out: list[str] = []
        for ev in _LD(report.get("evaluations")):
            for f in (ev.get("red_flags") or []):
                if isinstance(f, str):
                    out.append(f"_{ev.get('agent_role', 'Reviewer')}_: {f}")
            for w in (ev.get("weaknesses") or [])[:2]:
                if isinstance(w, str):
                    out.append(f"_{ev.get('agent_role', 'Reviewer')}_: {w}")
        gap_d = _D(report.get("gap_analysis"))
        for g in (gap_d.get("critical_gaps") or []):
            if isinstance(g, str):
                out.append(f"_Gap analysis_: {g}")
        # Dedup while preserving order.
        seen, deduped = set(), []
        for s in out:
            k = s.strip().lower()
            if k and k not in seen:
                seen.add(k)
                deduped.append(s)
        return deduped[:12]

    # ──────────────────────────────────────────────────────────────────────
    # Per the user's brief: tables out, bullets in; emojis ONLY on `##`
    # section headers (not inline, not on impact labels). Sections aim to
    # be concise — one beat per bullet, nest only when it adds context.
    # ──────────────────────────────────────────────────────────────────────
    badge = enrichment.get("status_badge") or ("", "Unassessed", "")
    badge_label = badge[1] if len(badge) > 1 else "Unassessed"
    badge_note = badge[2] if len(badge) > 2 else ""

    md.append("# Resume Recommendation Report")
    md.append(f"## {job.get('title') or 'Role'} @ {job.get('company') or 'Company'}\n")

    # ── §1.1 — TL;DR decision card (5-line, the only thing many users read) ─
    tldr = _D(enrichment.get("tldr"))
    if tldr:
        md.append("> **TL;DR — should you apply?**")
        md.append(">")
        md.append(f"> - **Verdict:** {tldr.get('verdict','—')}"
                   + (f" — _{tldr['verdict_reason']}_" if tldr.get("verdict_reason") else ""))
        eff_band = tldr.get("effort_band", "—")
        eff_label = tldr.get("effort_label", "—")
        eff_emoji = tldr.get("effort_emoji", "")
        md.append(f"> - **Effort to submit:** {eff_emoji} {eff_band} ({eff_label})")
        fresh = tldr.get("freshness") or "—"
        if fresh != "—":
            md.append(f"> - **Freshness:** {fresh}")
        if tldr.get("top_blocker"):
            md.append(f"> - **Top blocker:** {tldr['top_blocker']}")
        if tldr.get("top_lever"):
            md.append(f"> - **Top lever:** {tldr['top_lever']}")
        md.append("")

    # ── 📌 Job Posting Metadata (parsed from JD top-of-file block) ────────
    # Renders the canonical metadata block from each JD:
    #   Position, Link, Contract type, Remote, Company
    # (Status / Freshness bullets removed per user request — they're already
    # in the TL;DR card above.)
    jd_meta = _D(job.get("jd_metadata"))
    source_url = (job.get("source_url") or "").strip()
    if not source_url and job.get("source_file"):
        try:
            jd_path_str = job["source_file"]
            jd_p = Path(jd_path_str)
            if jd_p.exists() and jd_p.suffix.lower() in (".md", ".txt"):
                txt = jd_p.read_text(encoding="utf-8", errors="ignore")[:8000]
                m = re.search(r"(?im)^\**\s*link\s*\**\s*[:\-]\s*(\S+)", txt)
                if not m:
                    m = re.search(r"(?im)^(https?://\S+)", txt)
                if m:
                    source_url = m.group(1).strip().strip("<>")
        except Exception:
            pass

    # Prefer parsed metadata values; fall back to the agent's structured job.
    position = (jd_meta.get("position") or job.get("title") or "").strip()
    contract = (jd_meta.get("contract_type") or job.get("modality") or "").strip()
    remote_label = jd_meta.get("remote_label") or ""
    if not remote_label and jd_meta.get("location"):
        remote_label = f"📍 {jd_meta['location']}"
    if not remote_label and job.get("location"):
        remote_label = f"📍 {job['location']}"
    company_name = (jd_meta.get("company") or job.get("company") or "").strip()

    if position or source_url or contract or remote_label or company_name:
        md.append("## 📌 Job Posting Metadata\n")
        if position:
            md.append(f"- **Position:** {position}")
        if source_url:
            md.append(f"- **🔗 Link:** [{source_url}]({source_url})")
        if contract:
            md.append(f"- **Contract type:** {contract}")
        if remote_label:
            md.append(f"- **Remote:** {remote_label}")
        if company_name:
            md.append(f"- **Company:** {company_name}")
        if job.get("source_file"):
            md.append(f"- **Local file:** `{job['source_file']}`")
        md.append("")

    if report.get("executive_summary"):
        md.append(f"> {report['executive_summary']}\n")
    md.append("---\n")

    # ──────────────────────────────────────────────────────────────────────
    # Helper: build do / don't lists from the existing report data so the
    # candidate gets one clear, prescriptive section at the top — derived
    # from prioritized changes, JD verbs, ATS gaps, red flags, weak words.
    # ──────────────────────────────────────────────────────────────────────
    def _build_dos_donts() -> tuple:
        cf_d = _D(report.get("consolidated_feedback"))
        changes_d = _LD(cf_d.get("prioritized_changes"))
        ats_d = _D(report.get("ats_report"))
        gap_d = _D(report.get("gap_analysis"))
        weak_audit = enrichment.get("weak_words_audit") or []
        strong = enrichment.get("strong_verbs") or []

        dos: list[str] = []
        donts: list[str] = []

        # Do: top HIGH-impact prioritized changes.
        high = [c for c in changes_d if (c.get("impact") or "").lower() == "high"]
        for c in high[:5]:
            desc = (c.get("description") or "").strip()
            sect = (c.get("target_section") or "").strip()
            if desc:
                dos.append(f"{desc}" + (f" _(section: {sect})_" if sect else ""))

        # Do: mirror strong verbs from the JD.
        if strong:
            dos.append("Mirror the JD's strong verbs in your bullets: "
                       + ", ".join(f"`{v}`" for v in strong[:8]))

        # Do: inject the top missing ATS keywords.
        missing_kw = ats_d.get("missing_keywords") or []
        if missing_kw:
            dos.append("Weave these missing ATS terms into your summary and bullets: "
                       + ", ".join(f"`{k}`" for k in missing_kw[:10]))

        # Do: address critical gaps explicitly (acknowledge + reframe).
        for g in (gap_d.get("critical_gaps") or [])[:3]:
            if isinstance(g, str) and g.strip():
                dos.append(f"Address the critical gap up front: {g.strip()}")

        # Don't: weak / vague phrases the audit flagged.
        for w in weak_audit[:6]:
            phrases = ", ".join(f"`{p}`" for p in (w.get("weak_phrases") or [])[:3])
            if phrases:
                donts.append(f"Avoid weak phrasing — drop {phrases}")

        # Don't: overused keywords (variety beats repetition for ATS too).
        over = ats_d.get("overused_keywords") or []
        if over:
            donts.append("Don't keep repeating the same terms — vary: "
                         + ", ".join(f"`{k}`" for k in over[:6]))

        # Don't: red flags from reviewers (top ones).
        for rf in _collect_red_flags()[:5]:
            donts.append(f"Don't ignore this reviewer flag → {rf}")

        # Don't: inflate or invent — keep it grounded.
        donts.append("Don't claim metrics or scope you can't defend in the interview.")

        # Dedup while preserving order.
        def _dedup(items: list[str]) -> list[str]:
            seen, out = set(), []
            for s in items:
                k = s.strip().lower()
                if k and k not in seen:
                    seen.add(k)
                    out.append(s)
            return out

        return _dedup(dos)[:10], _dedup(donts)[:10]

    # ── Resume Recommendations — Must Do / Don't Do ──────────────────────
    dos, donts = _build_dos_donts()
    md.append("## ✅ Resume Recommendation — Must Do\n")
    if dos:
        for d in dos:
            md.append(f"- {d}")
        md.append("")
    else:
        md.append("- No high-impact actions detected — proofread and submit.\n")

    # §1.3 — The legacy "Don't Do" section was just paraphrased reviewer
    # flags. Replace with genuine anti-patterns: weak phrasing the audit
    # caught + ATS overuse + a single "don't inflate" reminder. Reviewer
    # flags get their own dedicated section below.
    md.append("## 🚫 Resume Recommendation — Don't Do\n")
    weak_audit = enrichment.get("weak_words_audit") or []
    ats_d_for_donts = _D(report.get("ats_report"))
    real_donts: list[str] = []
    for w in weak_audit[:5]:
        phrases = ", ".join(f"`{p}`" for p in (w.get("weak_phrases") or [])[:3])
        if phrases:
            real_donts.append(f"Avoid weak phrasing — drop {phrases}")
    over = ats_d_for_donts.get("overused_keywords") or []
    if over:
        real_donts.append("Don't keep repeating the same terms — vary: "
                          + ", ".join(f"`{k}`" for k in over[:6]))
    verification = _D(report.get("verification_report"))
    flagged_claims = verification.get("flagged_claims") or verification.get("issues") or []
    for fc in (flagged_claims if isinstance(flagged_claims, list) else [])[:3]:
        text = fc if isinstance(fc, str) else (fc.get("text") if isinstance(fc, dict) else "")
        if text:
            real_donts.append(f"Don't claim what you can't defend — verification flagged: {text}")
    real_donts.append("Don't claim metrics, dates, or scope you can't defend in the interview.")
    for d in real_donts:
        md.append(f"- {d}")
    md.append("")

    # ──────────────────────────────────────────────────────────────────────
    # Helper: append a section ONLY if it has substantive content. Sections
    # with no information are omitted entirely (no header, no placeholder),
    # per the user's brief. Sections whose "empty state" is itself useful
    # information (e.g. "Every bullet has a metric") opt into rendering by
    # passing a non-empty `body` themselves.
    # ──────────────────────────────────────────────────────────────────────
    def _section(title: str, body: list[str]) -> None:
        # Strip trailing blanks so a body of [""] doesn't count as content.
        meaningful = [line for line in body if line and line.strip()]
        if not meaningful:
            return
        md.append(f"## {title}\n")
        md.extend(body)
        if md and md[-1] != "":
            md.append("")

    # ── 📊 Match by Category ──────────────────────────────────────────────
    qmatch = enrichment.get("qualitative_match") or []
    body: list[str] = []
    for row in qmatch:
        body.append(f"- **{row.get('category','—')}** — {row.get('band','—')}")
        sig = (row.get("signal") or "").strip()
        if sig:
            body.append(f"  - {sig}")
    _section("📊 Match by Category", body)

    # ── 🤔 Should You Apply? ──────────────────────────────────────────────
    decision = enrichment.get("decision_helper") or {}
    body = []
    if decision:
        body.append(f"- **Verdict:** {decision.get('verdict', '—')}")
        reason = (decision.get("reason") or "").strip()
        if reason:
            body.append(f"- **Why:** {reason}")
        if decision.get("pros"):
            body.append("- **Pros:**")
            for p in decision["pros"]:
                body.append(f"  - {p}")
        if decision.get("cons"):
            body.append("- **Cons:**")
            for c in decision["cons"]:
                body.append(f"  - {c}")
    _section("🤔 Should You Apply?", body)

    # ── 🚩 Blockers & Gaps (deduped) ──────────────────────────────────────
    # §1.2 — single canonical list, each entry annotated with the reviewers
    # that flagged it. Replaces the redundant top-level Red Flags + reviewer
    # red_flag lists + duplicated Gap Analysis prose.
    body = []
    if canonical_blockers:
        body.append(f"- _Each blocker is shown once. \"Flagged by:\" tells you "
                    f"how many reviewers raised it (of {max(len(reviewer_roles),1)})._")
        for b in canonical_blockers:
            flaggers = ", ".join(b["flagged_by"])
            n = len(b["flagged_by"])
            md_kind = "🟥" if n >= 2 else ("🟧" if b["kind"] == "gap" else "🟨")
            body.append(f"- {md_kind} **{b['text']}**")
            body.append(f"  - _Flagged by:_ {flaggers}")
    else:
        body = ["- None — profile is aligned on the major fit dimensions."]
    _section("🚩 Blockers & Gaps", body)

    # ── ✏️ Pre-Submission Checklist ──────────────────────────────────────
    checklist = enrichment.get("pre_submission_checklist") or []
    body = [f"- [ ] {item}" for item in checklist]
    _section("✏️ Pre-Submission Checklist", body)

    # ── 🧾 Proposal Analysis & Plan ───────────────────────────────────────
    # Rendered only when the Opus-driven cv_proposal.docx has been generated
    # alongside this report. Describes what landed in the proposal and the
    # remaining manual steps before the candidate submits.
    proposal_analysis = _D(enrichment.get("proposal_analysis"))
    if proposal_analysis:
        body = []
        docx_path = proposal_analysis.get("docx_path") or ""
        if docx_path:
            from pathlib import Path as _P
            body.append(f"- **File:** `{_P(docx_path).name}` — open `cv_proposal.docx` next to this report.")
        highlights = proposal_analysis.get("highlights") or []
        if highlights:
            body.append("- **What's in the proposal:**")
            for h in highlights:
                body.append(f"  - {h}")
        diff = proposal_analysis.get("diff") or []
        if diff:
            body.append("- **Key changes vs. your current CV:**")
            for d in diff:
                if not isinstance(d, dict):
                    continue
                field = d.get("field", "")
                before = d.get("before", "")
                after = d.get("after", "")
                body.append(f"  - **{field}**")
                body.append(f"    - Before: {before}")
                body.append(f"    - After: {after}")
        covered = proposal_analysis.get("keywords_covered") or []
        if covered:
            body.append("- **ATS keywords woven in:** "
                        + ", ".join(f"`{k}`" for k in covered[:20]))
        still_missing = proposal_analysis.get("keywords_still_missing") or []
        if still_missing:
            body.append("- **Keywords still missing (acquire evidence before claiming):** "
                        + ", ".join(f"`{k}`" for k in still_missing[:15]))
        action_plan = proposal_analysis.get("action_plan") or []
        if action_plan:
            body.append("- **Action plan — finish these before submitting:**")
            by_pri: dict[str, list[str]] = {"High": [], "Medium": [], "Low": []}
            for ap in action_plan:
                if not isinstance(ap, dict):
                    continue
                pri = (ap.get("priority") or "Medium").title()
                by_pri.setdefault(pri, []).append(ap.get("step", ""))
            for pri in ("High", "Medium", "Low"):
                steps = [s for s in by_pri.get(pri, []) if s]
                if not steps:
                    continue
                body.append(f"  - **{pri} priority**")
                for s in steps:
                    body.append(f"    - [ ] {s}")
        _section("🧾 Proposal Analysis & Plan", body)

    # ── 📋 Recommended Resume Edits ───────────────────────────────────────
    cf = _D(report.get("consolidated_feedback"))
    changes = _LD(cf.get("prioritized_changes"))
    body = []
    if changes:
        by_impact = {"high": [], "medium": [], "low": []}
        for ch in changes:
            by_impact.setdefault((ch.get("impact") or "medium").lower(), []).append(ch)
        for level, label in (("high", "High impact"), ("medium", "Medium impact"), ("low", "Low impact")):
            if not by_impact.get(level):
                continue
            body.append(f"- **{label}**")
            for ch in by_impact[level]:
                ctype = (ch.get("change_type", "edit") or "edit").lower()
                section = ch.get("target_section") or "summary"
                effort = ch.get("effort", "?")
                desc = ch.get("description") or ""
                body.append(f"  - _{section}_ ({ctype}, effort {effort}) — {desc}")
    _section("📋 Recommended Resume Edits", body)

    # ── ✍️ Bullet-by-Bullet Rewrites (word-level diff) — §1.4 ─────────────
    rewrites = strategic.get("bullet_rewrites") or []
    body = []
    for br in rewrites:
        t = br.get("title") or "Role"
        co = br.get("company") or "Company"
        before = br.get("original", "") or ""
        after = br.get("rewritten", "") or ""
        body.append(f"- **{t} — {co}**")
        diff_line = _wordlevel_diff(before, after)
        body.append(f"  - {diff_line}")
        if br.get("reason"):
            body.append(f"  - _Why:_ {br['reason']}")
    _section("✍️ Bullet Rewrites", body)

    # ── 🎯 ATS Keywords ───────────────────────────────────────────────────
    ats = _D(report.get("ats_report"))
    body = []
    if ats.get("missing_keywords"):
        kws = ats["missing_keywords"][:25]
        body.append("- **Missing — weave naturally into bullets, summary, and skills:**")
        body.append(f"  - {', '.join(f'`{k}`' for k in kws)}")
    if ats.get("overused_keywords"):
        body.append("- **Overused (vary the language):**")
        body.append(f"  - {', '.join(f'`{k}`' for k in ats['overused_keywords'][:10])}")
    if ats.get("format_issues"):
        body.append("- **Format issues:**")
        for f in ats["format_issues"]:
            body.append(f"  - {f}")
    _section("🎯 ATS Keywords to Inject", body)

    # ── 🗺️ Skill Alignment ────────────────────────────────────────────────
    matrix = enrichment.get("skill_alignment_matrix") or []
    body = []
    if matrix:
        status_label = {"✅": "Full match", "⚠️": "Partial", "❌": "Missing"}
        groups: dict[str, list[dict]] = {}
        for row in matrix:
            label = status_label.get(row.get("status", ""), row.get("status", "Other"))
            groups.setdefault(label, []).append(row)
        for label in ("Full match", "Partial", "Missing"):
            if label not in groups:
                continue
            body.append(f"- **{label}**")
            for row in groups[label]:
                pri = row.get("priority", "")
                cat = row.get("category", "")
                req = row.get("requirement", "")
                body.append(f"  - _{pri}/{cat}_ — {req}")
    _section("🗺️ Skill Alignment", body)

    # ── 🔢 Quantification Audit ───────────────────────────────────────────
    quant = enrichment.get("quantification_audit") or []
    body = []
    if quant:
        body.append("- Bullets missing metrics — add a number where you can:")
        for q in quant:
            body.append(f"  - **{q.get('title','')}, {q.get('company','')}** — {q.get('bullet','')}")
    elif (candidate.get("work_experience") or []):
        # Positive empty state — only meaningful if there ARE bullets to audit.
        body = ["- Every bullet already has a metric."]
    _section("🔢 Quantification Audit", body)

    # ── 💪 Strong Verbs to Mirror ─────────────────────────────────────────
    strong = enrichment.get("strong_verbs") or []
    body = []
    if strong:
        body.append("- Mirror these verbs from the JD:")
        body.append(f"  - {' · '.join(f'`{v}`' for v in strong)}")
    _section("💪 Strong Verbs to Mirror", body)

    # ── 🪶 Weak Phrases to Replace ────────────────────────────────────────
    weak = enrichment.get("weak_words_audit") or []
    body = []
    if weak:
        for w in weak:
            body.append(f"- _{w.get('bullet','')}_")
            wp = ", ".join(f"`{p}`" for p in (w.get("weak_phrases") or []))
            sg = ", ".join(f"`{s}`" for s in (w.get("suggestions") or []))
            body.append(f"  - Replace {wp} → {sg}")
    elif (candidate.get("work_experience") or []):
        body = ["- No weak phrasing detected."]
    _section("🪶 Weak Phrases to Replace", body)

    # ── 🧭 Career Narrative ───────────────────────────────────────────────
    narrative = strategic.get("career_narrative")
    _section("🧭 Career Narrative", [f"> {narrative}"] if narrative else [])

    # ── 🤝 Cultural Fit ───────────────────────────────────────────────────
    culture = strategic.get("cultural_fit_signals")
    body = []
    if culture:
        body.append(f"> {culture}")
    if job.get("cultural_signals"):
        body.append("- **Signals from JD language:**")
        for s in job["cultural_signals"]:
            body.append(f"  - {s}")
    _section("🤝 Cultural Fit", body)

    # ── 📜 Cover Letter — outline + body preview ──────────────────────────
    # §1.6 — render the actual cover letter body when the pipeline produced
    # one (cover_letter_task), in addition to the strategic outline. The
    # full version is also written as cover_letter.pdf next to this report.
    cl_outline = strategic.get("cover_letter_outline") or []
    cl_body_obj = _D(report.get("cover_letter"))
    cl_body_text = (cl_body_obj.get("body") or cl_body_obj.get("text") or "").strip()
    body = []
    if cl_outline:
        body.append("- **Outline (structural guidance):**")
        for line in cl_outline:
            body.append(f"  - {line}")
    if cl_body_text:
        body.append("")
        body.append("<details><summary>📝 Show generated cover letter draft "
                    "(also saved as `cover_letter.pdf`)</summary>")
        body.append("")
        for line in cl_body_text.splitlines():
            line = line.strip()
            if line:
                body.append(f"> {line}")
            else:
                body.append(">")
        body.append("")
        body.append("</details>")
        body.append("")
    _section("📜 Cover Letter", body)

    # ── 🎯 Application Strategy ───────────────────────────────────────────
    strategy = strategic.get("application_strategy")
    _section("🎯 Application Strategy", [f"> {strategy}"] if strategy else [])

    # ── 🎤 Likely Interview Questions ─────────────────────────────────────
    iq = _LD(report.get("interview_questions"))
    body = []
    for q in iq:
        body.append(f"- **{q.get('question','')}**")
        if q.get("reasoning"):
            body.append(f"  - Why: {q['reasoning']}")
        if q.get("suggested_angle"):
            body.append(f"  - Angle: {q['suggested_angle']}")
    _section("🎤 Likely Interview Questions", body)

    # ── ❓ Questions YOU Should Ask ────────────────────────────────────────
    qta = strategic.get("questions_to_ask_interviewer") or []
    _section("❓ Questions YOU Should Ask", [f"- {q}" for q in qta])

    # ── 🎒 What to Bring to the Interview ─────────────────────────────────
    prep = strategic.get("interview_preparation") or []
    _section("🎒 What to Bring to the Interview", [f"- [ ] {p}" for p in prep])

    # ── 🛠️ Skills to Develop ──────────────────────────────────────────────
    skills_dev = strategic.get("skills_to_develop") or []
    _section("🛠️ Skills to Develop", [f"- {s}" for s in skills_dev])

    # ── 🔍 Gap Analysis (minor + framing only — criticals live in Blockers) ─
    # §1.2 — critical_gaps moved to the canonical Blockers section above.
    gap = _D(report.get("gap_analysis"))
    body = []
    if gap.get("minor_gaps"):
        body.append("- **Minor** (nice-to-have):")
        for g in gap["minor_gaps"]:
            body.append(f"  - {g}")
    if gap.get("framing_opportunities"):
        body.append("- **Framing opportunities:**")
        for g in gap["framing_opportunities"]:
            body.append(f"  - {g}")
    _section("🔍 Gap Analysis", body)

    # ── 👥 Reviewer Insights — comparative grid + collapsed transcripts ────
    # §1.5 — the ~200-line per-reviewer dump is collapsed into a single
    # comparative table (Strengths/Weaknesses/Red flags counts per reviewer)
    # plus a <details> block with the full transcripts for the rare user
    # who needs forensic depth.
    body = []
    evs = _LD(report.get("evaluations"))
    if evs:
        # Header row
        cols = ["Reviewer", "Strengths", "Weaknesses", "Red flags", "Recommendations"]
        body.append("| " + " | ".join(cols) + " |")
        body.append("|" + "|".join(["---"] * len(cols)) + "|")
        for ev in evs:
            role = ev.get("agent_role") or "—"
            s = len(ev.get("strengths") or [])
            w = len(ev.get("weaknesses") or [])
            rf = len(ev.get("red_flags") or [])
            recs = len(ev.get("recommendations") or [])
            body.append(f"| **{role}** | {s} | {w} | {rf} | {recs} |")
        body.append("")
        body.append("<details><summary>📜 Show full reviewer transcripts</summary>")
        body.append("")
        for ev in evs:
            role = ev.get("agent_role") or "—"
            body.append(f"#### {role}\n")
            for field_key, label in (
                ("strengths", "Strengths"),
                ("weaknesses", "Weaknesses"),
                ("recommendations", "Recommendations"),
                ("uncertain_items", "Verify yourself"),
            ):
                items = ev.get(field_key) or []
                if items:
                    body.append(f"- **{label}:**")
                    for v in items:
                        body.append(f"  - {v}")
            body.append("")
        body.append("</details>")
        body.append("")
    _section("👥 Reviewer Insights", body)

    so = _D(report.get("second_opinion"))
    if so and so.get("triggered") and so.get("rationale"):
        _section("⚖️ Second Opinion", [f"> {so.get('rationale','')}"])

    # ── 💰 Compensation Reference — §1.7 ──────────────────────────────────
    body = []
    if job.get("salary_range"):
        body.append(f"- **Posting range:** {job['salary_range']}")
    _section("💰 Compensation Reference", body)

    # ── 🥊 Hypothetical Competitor ────────────────────────────────────────
    cp = _D(report.get("competitor_profile"))
    body = []
    if cp:
        if cp.get("summary"):
            body.append(f"- _{cp['summary']}_")
        if cp.get("differentiators"):
            body.append("- **They have, you don't:**")
            for d in cp["differentiators"]:
                body.append(f"  - {d}")
        if cp.get("candidate_advantages"):
            body.append("- **You have, they don't:**")
            for a in cp["candidate_advantages"]:
                body.append(f"  - {a}")
    _section("🥊 Hypothetical Competitor", body)

    # ── 🏢 Company Snapshot ───────────────────────────────────────────────
    snapshot = strategic.get("company_snapshot")
    _section("🏢 Company Snapshot", [f"> {snapshot}"] if snapshot else [])

    # ── 🇨🇷 Company in Costa Rica ─────────────────────────────────────────
    # Haiku-generated paragraph describing the company's known CR
    # operations (office presence, business function, headcount when
    # publicly reported). Cached per-company so a 16-job run only pays
    # once per unique employer.
    cr_blurb = (enrichment.get("company_costa_rica") or "").strip()
    if cr_blurb:
        company_label = (job.get("company") or "the company").strip()
        body = [
            f"> _Public-information summary of {company_label}'s Costa "
            f"Rica operations. Verify before acting._",
            "",
            f"> {cr_blurb}",
        ]
        _section("🇨🇷 Company in Costa Rica", body)

    # ── 💼 Job Posting Reference ──────────────────────────────────────────
    body = []
    if job and any(job.get(k) for k in ("title", "company", "location", "modality",
                                          "seniority", "salary_range", "role_type",
                                          "tech_stack", "ats_keywords", "responsibilities",
                                          "requirements")):
        md.append("---\n")
        for k_label, v in (
            ("Title", job.get("title", "")),
            ("Company", job.get("company", "")),
            ("Location", job.get("location") or ""),
            ("Modality", job.get("modality") or ""),
            ("Seniority", job.get("seniority") or ""),
            ("Salary range", job.get("salary_range") or ""),
            ("Role type", job.get("role_type") or ""),
        ):
            if v:
                body.append(f"- **{k_label}:** {v}")
        if job.get("source_url"):
            body.append(f"- **Source URL:** [{job['source_url']}]({job['source_url']})")
        if job.get("tech_stack"):
            body.append("- **Tech stack:** " + ", ".join(job["tech_stack"]))
        if job.get("ats_keywords"):
            body.append("- **ATS keywords (job posting):** "
                        + ", ".join(job["ats_keywords"][:25]))
        if job.get("responsibilities"):
            body.append("- **Key responsibilities:**")
            for r in job["responsibilities"][:8]:
                body.append(f"  - {r}")
        reqs = job.get("requirements") or []
        if reqs:
            body.append("- **Requirements:**")
            for r in reqs[:12]:
                if isinstance(r, dict):
                    txt = r.get("text", "")
                    pri = r.get("priority", "")
                    cat = r.get("category", "")
                    body.append(f"  - _{pri}/{cat}_ — {txt}")
                else:
                    body.append(f"  - {r}")
    _section("💼 Job Posting Reference", body)

    # ── 👤 Parsed Resume — collapsed §1.8 ─────────────────────────────────
    # The full parsed-resume dump is debugging artifact for most readers. It
    # is now (a) collapsed behind <details> and (b) also written verbatim as
    # `parsed_resume.json` next to this report by the orchestrator. Keep the
    # markdown rendition concise.
    body = []
    body.append("<details><summary>👤 Show parsed resume (also saved as "
                "`parsed_resume.json`)</summary>\n")
    body.append("_Source of truth: the parsed CV the reviewers actually "
                "scored against. Verify before submitting._\n")
    if candidate.get("source_file"):
        body.append(f"- **Parsed from:** `{candidate['source_file']}`")

    name = candidate.get("full_name") or "_not parsed_"
    body.append(f"- **Name:** {name}")
    if candidate.get("headline"):
        body.append(f"- **Headline:** {candidate['headline']}")

    contact = _D(candidate.get("contact"))
    contact_rows = [(k, v) for k, v in contact.items() if v]
    if contact_rows:
        body.append("- **Contact:**")
        for k, v in contact_rows:
            label = k.replace("_", " ").title()
            body.append(f"  - **{label}:** {v}")
    else:
        body.append("- **Contact:** _not parsed_")

    if candidate.get("summary"):
        body.append(f"- **Summary:** {candidate['summary']}")
    else:
        body.append("- **Summary:** _not parsed_")

    def _fmt_skill(s):
        if isinstance(s, str):
            return s.strip()
        if isinstance(s, dict):
            return (s.get("name") or s.get("skill") or s.get("label") or "").strip()
        return ""
    skills = [v for v in (_fmt_skill(s) for s in (candidate.get("skills") or [])) if v]
    if skills:
        body.append(f"- **Skills ({len(skills)}):**")
        body.append("  - " + ", ".join(f"`{s}`" for s in skills))
    else:
        body.append("- **Skills:** _not parsed_")

    def _fmt_language(l):
        if isinstance(l, str):
            return l
        if isinstance(l, dict):
            name = l.get("language") or l.get("name") or ""
            prof = l.get("proficiency") or l.get("fluency") or l.get("level") or ""
            return f"{name} — {prof}" if (name and prof) else (name or prof)
        return ""
    languages = [s for s in (_fmt_language(l) for l in (candidate.get("languages") or [])) if s]
    if languages:
        body.append("- **Languages:** " + ", ".join(languages))

    def _fmt_cert(c):
        if isinstance(c, str):
            return c
        if isinstance(c, dict):
            name = c.get("certification_name") or c.get("name") or c.get("title") or ""
            org = c.get("issuing_organization") or c.get("issuer") or c.get("organization") or ""
            date = c.get("date_obtained") or c.get("date") or c.get("year") or ""
            parts = [p for p in [name, org, date] if p and p != "unknown"]
            return " — ".join(parts) if parts else ""
        return ""
    certs = [s for s in (_fmt_cert(c) for c in (candidate.get("certifications") or [])) if s]
    if certs:
        body.append("- **Certifications:**")
        for c in certs:
            body.append(f"  - {c}")

    wx = [w for w in (candidate.get("work_experience") or []) if isinstance(w, dict)]
    if wx:
        body.append(f"- **Work Experience ({len(wx)} role{'s' if len(wx) != 1 else ''}):**")
        for w in wx:
            t = w.get("title") or w.get("position") or w.get("role") or "—"
            co = w.get("company") or w.get("employer") or "—"
            start = w.get("start_date", "")
            end = w.get("end_date") or "Present"
            loc = w.get("location") or ""
            period = f"{start} – {end}" + (f", {loc}" if loc else "")
            body.append(f"  - **{t} — {co}** _{period}_")
            for b in (w.get("bullets") or w.get("achievements") or w.get("responsibilities") or []):
                if isinstance(b, str) and b.strip():
                    body.append(f"    - {b}")
                elif isinstance(b, dict):
                    txt = b.get("text") or b.get("bullet") or b.get("description") or ""
                    if txt:
                        body.append(f"    - {txt}")
            techs = w.get("technologies") or w.get("tools") or w.get("tech_stack")
            if techs:
                techs_str = ", ".join(t for t in techs if isinstance(t, str))
                if techs_str:
                    body.append(f"    - _Technologies: {techs_str}_")
    else:
        body.append("- **Work Experience:** _not parsed_")

    edu = [e for e in (candidate.get("education") or []) if isinstance(e, dict)]
    if edu:
        body.append("- **Education:**")
        for e in edu:
            deg = e.get("degree") or e.get("title") or ""
            field = e.get("field") or e.get("field_of_study") or ""
            inst = e.get("institution") or e.get("school") or e.get("university") or ""
            start = e.get("start_date") or ""
            end = e.get("end_date") or e.get("graduation_year") or ""
            period = f" ({start}–{end})" if start or end else ""
            label = deg + (f", {field}" if field else "")
            sep = " — " if label and inst else ""
            body.append(f"  - **{label or inst}**{sep}{inst if label else ''}{period}")
    else:
        body.append("- **Education:** _not parsed_")
    body.append("\n</details>")
    _section("👤 Parsed Resume", body)

    # ── 🔗 Useful Links ──────────────────────────────────────────────────
    links = enrichment.get("links") or []
    _section("🔗 Useful Links", [f"- [{link['label']}]({link['url']})" for link in links])

    # ── Footer ───────────────────────────────────────────────────────────
    md.append("---\n")
    md.append("_This report was generated by **CV Optimizer** — an automated "
              "multi-agent analysis. The recommendations are starting points; "
              "always tailor edits to your authentic experience._\n")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text("\n".join(md), encoding="utf-8")
    return json.dumps({"status": "ok", "output_path": output_path})


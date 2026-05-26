"""CV document parser tool.

Supports DOCX (via python-docx) and PDF (via pdfplumber).
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from crewai.tools import tool

__all__ = ["parse_cv", "parse_cv_docx"]


@tool("parse_cv")
def parse_cv(cv_path: str) -> str:
    """
    Reads a CV in DOCX or PDF and extracts its content. Auto-detects
    format by extension. Returns JSON with paragraphs (and tables for
    DOCX) plus a `source_format` field so downstream tools know whether
    they can use the file as a formatting template (DOCX) or must build
    a fresh DOCX from scratch (PDF).
    """
    path = Path(cv_path)
    if not path.exists():
        return json.dumps({"error": f"file not found: {cv_path}"})

    ext = path.suffix.lower()
    if ext == ".docx":
        return _parse_cv_docx_impl(path)
    if ext == ".pdf":
        return _parse_cv_pdf_impl(path)
    return json.dumps(
        {
            "error": f"unsupported CV format: {ext}. Supported: .docx, .pdf",
        }
    )


def _parse_cv_docx_impl(path: Path) -> str:
    """DOCX parser — preserves paragraph styles + tables for templating."""
    try:
        from docx import Document
    except ImportError:
        return json.dumps({"error": "python-docx missing: pip install python-docx"})

    try:
        doc = Document(path)
    except Exception as e:
        return json.dumps({"error": f"failed to open DOCX: {e}"})

    paragraphs: list[dict[str, Any]] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        paragraphs.append({"text": text, "style": p.style.name if p.style else "Normal"})

    tables_text: list[list[list[str]]] = []
    for tbl in doc.tables:
        tables_text.append([[cell.text.strip() for cell in row.cells] for row in tbl.rows])

    return json.dumps(
        {
            "source_file": str(path),
            "source_format": "docx",
            "can_use_as_template": True,
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables_text,
        }
    )


def _parse_cv_pdf_impl(path: Path) -> str:
    """PDF parser — extracts text per-line, infers paragraph boundaries."""
    try:
        import pdfplumber
    except ImportError:
        return json.dumps({"error": "pdfplumber missing: pip install pdfplumber"})

    paragraphs: list[dict[str, Any]] = []
    page_count = 0
    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                txt = page.extract_text() or ""
                # Each non-empty line becomes a paragraph. Bullet lines
                # are tagged with a "List Bullet" style hint so the
                # output renderer treats them correctly.
                for line in txt.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    is_bullet = bool(re.match(r"^[•\-\*▪◦·]\s", line))
                    style = "List Bullet" if is_bullet else "Normal"
                    if is_bullet:
                        line = re.sub(r"^[•\-\*▪◦·]\s+", "", line)
                    paragraphs.append({"text": line, "style": style})
    except Exception as e:
        return json.dumps({"error": f"failed to parse PDF CV: {e}"})

    return json.dumps(
        {
            "source_file": str(path),
            "source_format": "pdf",
            "can_use_as_template": False,  # PDFs cannot be reused as DOCX templates
            "page_count": page_count,
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "tables": [],
        }
    )


# Backwards-compatible alias so existing code paths keep working.
parse_cv_docx = parse_cv

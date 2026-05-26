"""Report rendering — Markdown/PDF output from a JobReport dict.

Contains all helpers for building and writing the final recommendation
report. Extracted from main.py (SRP Phase 3d). No logic changes.
"""

from __future__ import annotations

import contextlib
import json
import os
import re
import time
from pathlib import Path
from typing import Any

from src.cli.utils import slugify as _slugify
from src.company import resolve_company_domain as _resolve_company_domain
from src.pipeline.coercion import (
    as_dict as _as_dict,
)
from src.pipeline.coercion import (
    as_list as _as_list,
)
from src.pipeline.coercion import (
    as_list_of_dicts as _as_list_of_dicts,
)
from src.pipeline.scoring import (
    audit_quantification as _audit_quantification,
)
from src.pipeline.scoring import (
    audit_weak_words as _audit_weak_words,
)
from src.pipeline.scoring import (
    build_pre_submission_checklist as _build_pre_submission_checklist,
)
from src.pipeline.scoring import (
    compute_skill_alignment_matrix as _compute_skill_alignment_matrix,
)
from src.pipeline.scoring import (
    effort_impact_quadrant as _effort_impact_quadrant,
)
from src.pipeline.scoring import (
    extract_strong_verbs_from_jd as _extract_strong_verbs_from_jd,
)
from src.pipeline.scoring import (
    status_badge as _status_badge,
)
from src.presentation import (
    HAS_RICH as _HAS_RICH,
)
from src.presentation import (
    console as _console,
)
from src.presentation import (
    info as _info,
)
from src.presentation import (
    warn as _warn,
)
from src.report.enricher import (
    classify_match_with_opus as _classify_match_with_opus,
)
from src.report.enricher import (
    copy_job_source_as_pdf as _copy_job_source_as_pdf,
)
from src.report.enricher import (
    decision_helper_with_opus as _decision_helper_with_opus,
)
from src.report.enricher import (
    posting_freshness as _posting_freshness,
)

__all__: list[str] = []


def _build_links_section(report: dict, job: dict) -> list[dict]:
    """Aggregate every clickable URL into one section. §1.9 — favours
    company-specific intelligence (career page, recruiter LinkedIn from
    JD body, exact-name LinkedIn people search) over generic catch-all
    searches."""
    links: list[dict] = []
    company = (job.get("company") or "").strip()
    title = (job.get("title") or "").strip()
    location = (job.get("location") or "").strip()
    source_url = (job.get("source_url") or "").strip()
    source_file = (job.get("source_file") or "").strip()

    # ── Canonical posting links first ─────────────────────────────────────
    if source_url:
        links.append({"label": "Original job posting (source URL)", "url": source_url})
    if source_file:
        links.append({"label": "Original job posting (local file)", "url": f"file://{source_file}"})

    # ── Recruiter / hiring-manager handles parsed from JD body ────────────
    # Pull email + LinkedIn URL from the JD body if present.
    jd_body = ""
    if source_file:
        try:
            p = Path(source_file)
            if p.exists() and p.suffix.lower() in (".md", ".txt"):
                jd_body = p.read_text(encoding="utf-8", errors="ignore")[:8000]
        except Exception:
            pass
    if jd_body:
        for email_m in re.finditer(r"[\w.+-]+@[\w-]+\.[\w.-]+", jd_body):
            email = email_m.group(0)
            # Skip generic noreply / careers boilerplate.
            if any(x in email.lower() for x in ("noreply", "no-reply", "donotreply")):
                continue
            subj = f"Re: {title}" if title else "Application"
            links.append(
                {
                    "label": f"Email recruiter — {email}",
                    "url": f"mailto:{email}?subject={subj}",
                }
            )
            break
        for li_m in re.finditer(
            r"https?://(?:[a-z]{2,3}\.)?linkedin\.com/in/[\w\-_/]+", jd_body, re.I
        ):
            links.append({"label": "Recruiter LinkedIn (from JD)", "url": li_m.group(0)})
            break

    if company:
        from urllib.parse import quote_plus

        c = quote_plus(company)
        t = quote_plus(title)
        loc = quote_plus(location)
        # Try a company careers page guess from the resolver used by `search`.
        try:
            domain = _resolve_company_domain(company)
        except Exception:
            domain = None
        if domain:
            links.append(
                {"label": f"{company} careers page (guess)", "url": f"https://{domain}/careers"}
            )
            links.append({"label": f"{company} website", "url": f"https://{domain}"})

        # Company-specific LinkedIn jobs page (the brand, not the catch-all).
        links.append(
            {
                "label": f"LinkedIn jobs at {company}",
                "url": f"https://www.linkedin.com/jobs/search/?keywords={c}&location={loc}",
            }
        )
        # People search anchored on title + company (more useful than just company).
        if title:
            links.append(
                {
                    "label": f"Find {title} contacts at {company} (LinkedIn)",
                    "url": f"https://www.linkedin.com/search/results/people/"
                    f"?keywords={c}%20{t}&origin=GLOBAL_SEARCH_HEADER",
                }
            )
        else:
            links.append(
                {
                    "label": f"Find people at {company} (LinkedIn)",
                    "url": f"https://www.linkedin.com/search/results/people/"
                    f"?keywords={c}&origin=GLOBAL_SEARCH_HEADER",
                }
            )
        links.append(
            {
                "label": f"Glassdoor reviews — {company}",
                "url": f"https://www.glassdoor.com/Search/results.htm?keyword={c}",
            }
        )
    return links


# ─── Rich UX helpers for the strategic / proposal stages ─────────────────

# Each entry lists the inputs we hand to the LLM and the deliverables we
# expect back. Drives both the "about to call" preview panel and the
# post-call "what we got" summary table.
_STRATEGIC_DELIVERABLES = [
    ("✍️", "Bullet rewrites", "bullet_rewrites"),
    ("🧭", "Career narrative", "career_narrative"),
    ("🤝", "Cultural fit signals", "cultural_fit_signals"),
    ("📜", "Cover letter outline", "cover_letter_outline"),
    ("🎯", "Application strategy", "application_strategy"),
    ("❓", "Questions to ask", "questions_to_ask_interviewer"),
    ("🛠️", "Skills to develop", "skills_to_develop"),
    ("🎒", "Interview prep checklist", "interview_preparation"),
    ("🏢", "Company snapshot", "company_snapshot"),
]

_PROPOSAL_DELIVERABLES = [
    ("🪪", "Headline", "headline"),
    ("📝", "Professional summary", "summary"),
    ("🛠️", "Skills section", "skills"),
    ("💼", "Experience bullets", "experience"),
    ("🎓", "Education", "education"),
    ("🏷️", "ATS keywords woven in", "keywords_injected"),
    ("📋", "Notes for candidate", "notes_for_candidate"),
]


def _stage_panel(
    title: str, subtitle: str, items: list[tuple], model: str, border: str = "cyan"
) -> None:
    """Render a Rich panel announcing a long-running LLM stage with a
    preview of every artifact the call is about to produce. No-op when
    Rich is unavailable — falls back to a single plain line."""
    if not _HAS_RICH:
        print(f"\n→ {title} ({model})")
        for emoji, label, _key in items:
            print(f"   {emoji} {label}")
        return
    from rich import box as _box
    from rich.panel import Panel as _Panel
    from rich.table import Table as _Table
    from rich.text import Text as _Text

    grid = _Table.grid(padding=(0, 2))
    grid.add_column(justify="right", width=3)
    grid.add_column(justify="left")
    for emoji, label, _key in items:
        grid.add_row(emoji, _Text(label, style="white"))

    header = _Text()
    header.append(title, style=f"bold {border}")
    header.append("  •  ", style="dim")
    header.append(model, style="bold magenta")

    _console.print(
        _Panel(
            grid,
            title=header,
            subtitle=f"[dim]{subtitle}[/dim]",
            border_style=border,
            box=_box.ROUNDED,
            padding=(1, 2),
        )
    )


def _stage_status(message: str, spinner: str = "dots"):
    """Context-manager that shows a Rich spinner during the LLM call.
    Returns a no-op manager when Rich is unavailable."""
    if not _HAS_RICH:

        class _Null:
            def __enter__(self):
                print(f"   … {message}")
                return self

            def __exit__(self, *a):
                return False

        return _Null()
    return _console.status(
        f"[bold cyan]{message}[/bold cyan]", spinner=spinner, spinner_style="cyan"
    )


def _summarize_stage_result(
    title: str, result: dict, items: list[tuple], elapsed: float, model: str, border: str = "green"
) -> None:
    """Render a Rich table summarising which deliverables the LLM actually
    returned, with an item count per section."""
    if not _HAS_RICH:
        ok = sum(1 for _, _, k in items if (result or {}).get(k))
        print(f"   ✓ {title}: {ok}/{len(items)} sections in {elapsed:.1f}s ({model})")
        return
    from rich import box as _box
    from rich.table import Table as _Table

    table = _Table(
        title=f"[bold]{title}[/bold]  [dim]· {elapsed:.1f}s · {model}[/dim]",
        box=_box.ROUNDED,
        border_style=border,
        show_lines=False,
        header_style="bold",
        title_justify="left",
    )
    table.add_column("", width=3)
    table.add_column("Section", style="white", no_wrap=True)
    table.add_column("Status", justify="left")
    table.add_column("Items", justify="right", style="cyan")

    for emoji, label, key in items:
        val = (result or {}).get(key)
        if isinstance(val, list):
            count = len(val)
            status = "[bold green]✓ generated[/bold green]" if count else "[yellow]∅ empty[/yellow]"
            count_str = str(count) if count else "—"
        elif isinstance(val, str):
            count = len(val.split()) if val.strip() else 0
            status = "[bold green]✓ generated[/bold green]" if count else "[yellow]∅ empty[/yellow]"
            count_str = f"{count} words" if count else "—"
        elif isinstance(val, dict):
            count = len(val)
            status = "[bold green]✓ generated[/bold green]" if count else "[yellow]∅ empty[/yellow]"
            count_str = f"{count} keys" if count else "—"
        else:
            status = "[red]✗ missing[/red]"
            count_str = "—"
        table.add_row(emoji, label, status, count_str)
    _console.print(table)


def _generate_strategic_insights(report: dict, *, client: Any = None) -> dict:
    """One bundled Haiku call that produces all the medium/strategic
    LLM-derived sections in a single round-trip. Returns {} on any failure
    so the report falls back to the easy-wins-only version.
    """
    if client is None:
        try:
            from src.llm.client import get_default_client

            client = get_default_client()
        except Exception:
            return {}
    _llm = client

    job = _as_dict(report.get("job"))
    cand = _as_dict(report.get("candidate_profile"))
    cf = _as_dict(report.get("consolidated_feedback"))
    gap = _as_dict(report.get("gap_analysis"))

    # Pick the top 5 weakest bullets to suggest rewrites for.
    bullets_to_rewrite: list[dict] = []
    for w in _as_list_of_dicts(cand.get("work_experience"))[:3]:
        for b in _as_list(w.get("bullets"))[:3]:
            if b:
                bullets_to_rewrite.append(
                    {
                        "title": w.get("title", ""),
                        "company": w.get("company", ""),
                        "bullet": b[:280],
                    }
                )
            if len(bullets_to_rewrite) >= 5:
                break
        if len(bullets_to_rewrite) >= 5:
            break

    prompt = f"""You are a career strategist producing a single JSON object with
strategic insights for a candidate applying to this role. Return ONLY valid JSON
matching the schema below — no preamble, no markdown fences.

JOB POSTING:
- Title: {job.get("title", "")}
- Company: {job.get("company", "")}
- Location: {job.get("location", "")}
- Seniority: {job.get("seniority", "")}
- Modality: {job.get("modality", "")}
- Key responsibilities: {_as_list(job.get("responsibilities"))[:6]}
- Top requirements: {[r.get("text", "") if isinstance(r, dict) else r for r in _as_list(job.get("requirements"))[:8]]}
- Cultural signals: {_as_list(job.get("cultural_signals"))}

CANDIDATE:
- Name: {cand.get("full_name", "")}
- Headline: {cand.get("headline", "")}
- Top skills: {_as_list(cand.get("skills"))[:15]}
- Certifications: {_as_list(cand.get("certifications"))[:5]}

EVALUATION RESULT:
- Match score: {report.get("overall_match_score", 0)}/100
- Critical gaps: {_as_list(gap.get("critical_gaps"))[:5]}
- Top prioritized changes: {[c.get("description", "") for c in _as_list(cf.get("prioritized_changes"))[:5]]}

BULLETS TO REWRITE (return one per input):
{json.dumps(bullets_to_rewrite, indent=2)}

Return JSON with these keys (all required):
{{
  "bullet_rewrites": [
    {{"original": "...", "rewritten": "...", "reason": "...", "title": "...", "company": "..."}}
  ],
  "cover_letter_outline": ["bullet 1", "bullet 2", "bullet 3", "bullet 4", "bullet 5"],
  "career_narrative": "One paragraph: how the candidate should frame their career story for THIS role. 3-4 sentences max.",
  "cultural_fit_signals": "One paragraph: what the JD's tone/cultural signals say about the company, and how the candidate should mirror that in tone. 3-4 sentences.",
  "questions_to_ask_interviewer": ["question 1", "question 2", "question 3", "question 4", "question 5"],
  "skills_to_develop": ["skill 1 (rationale)", "skill 2 (rationale)", "skill 3 (rationale)"],
  "application_strategy": "One paragraph: best way to apply, when, who to contact, referral suggestions. 3-4 sentences.",
  "company_snapshot": "One paragraph based ONLY on what the JD reveals about the company stage/industry/mission. Do NOT invent facts. 3-4 sentences.",
  "interview_preparation": ["prep item 1", "prep item 2", "prep item 3", "prep item 4"]
}}"""

    from src.settings import get_settings as _gs

    try:
        content = _llm.complete(
            model=_gs().model_haiku,
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}],
        ).strip()
        if content.startswith("```"):
            content = re.sub(r"^```[a-z]*\n?", "", content)
            content = re.sub(r"\n?```$", "", content)
        return json.loads(content)
    except Exception as e:
        _warn(f"Strategic-insights LLM call failed: {e}")
        return {}


def _extract_job_link_from_file(job_path: Path) -> str | None:
    """Pull the JOB Link from a .md/.txt JD file. Conventions supported:
        Link: <url>            (explicit pattern)
        **Link:** <url>        (markdown-bold variant)
        <url>                  (bare URL line near the top — fallback)
    Returns None for PDF/DOCX or when no URL is found."""
    try:
        if not job_path.exists() or job_path.suffix.lower() not in (".md", ".txt"):
            return None
        text = job_path.read_text(encoding="utf-8", errors="ignore")[:8000]
    except Exception:
        return None
    m = re.search(r"(?im)^\**\s*link\s*\**\s*[:\-]\s*(\S+)", text)
    if m:
        return m.group(1).strip().strip("<>")
    for line in text.splitlines()[:25]:
        m = re.match(r"\s*(https?://\S+)\s*$", line)
        if m:
            return m.group(1).strip().strip("<>")
    return None


def _extract_jd_metadata_from_file(job_path: Path) -> dict[str, str]:
    """Parse the top-of-file metadata block emitted by our JD scrapers.

    Supported fields (case-insensitive, allows `**Bold:**` markdown):
        Position: <title>     — falls back to first non-empty line if absent
        Link: <url>
        Contract type: <…>    (or `Contract: <…>` for Workday-style files)
        Remote: <Yes|No|Si|…> (or `Location: <…>` when remote isn't stated)
        Company: <name>

    Returns a dict with present keys only (so the renderer can skip blanks).
    """
    md: dict[str, str] = {}
    if not job_path.exists() or job_path.suffix.lower() not in (".md", ".txt"):
        return md
    try:
        text = job_path.read_text(encoding="utf-8", errors="ignore")[:8000]
    except Exception:
        return md

    for key, pattern in _JD_METADATA_PATTERNS:
        m = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if m:
            md[key] = m.group(1).strip()

    # If Position wasn't explicit, fall back to the first non-empty line
    # that isn't itself a metadata line.
    if "position" not in md:
        for line in text.splitlines()[:8]:
            s = line.strip()
            if not s:
                continue
            if re.match(
                r"^\**\s*(position|link|contract|remote|location|company)\s*\**\s*[:\-]",
                s,
                re.IGNORECASE,
            ):
                continue
            if s.startswith("http://") or s.startswith("https://"):
                continue
            # Strip leading markdown heading hashes.
            md["position"] = re.sub(r"^#+\s*", "", s).strip()
            break

    # Normalize the remote flag to a friendly label.
    if "remote" in md:
        v = md["remote"].lower()
        if v in ("yes", "si", "sí", "true", "y", "remote"):
            md["remote_label"] = "🟢 Remote"
        elif v in ("hybrid", "híbrido", "hibrido"):
            md["remote_label"] = "🟡 Hybrid"
        elif v in ("no", "n", "on-site", "onsite", "on_site", "office"):
            md["remote_label"] = "🟠 On-site"
        else:
            md["remote_label"] = md["remote"]

    return md


# Field-name → regex mapping for the top-of-file JD metadata block.
# Order matters: first match per key wins. Adding a new field = adding one entry here.
_JD_METADATA_PATTERNS: list[tuple[str, str]] = [
    ("position", r"^\**\s*position\s*\**\s*[:\-]\s*(.+?)\s*$"),
    ("link", r"^\**\s*link\s*\**\s*[:\-]\s*(\S+).*$"),
    ("contract_type", r"^\**\s*contract(?:\s*type)?\s*\**\s*[:\-]\s*(.+?)\s*$"),
    ("remote", r"^\**\s*remote\s*\**\s*[:\-]\s*(.+?)\s*$"),
    ("location", r"^\**\s*location\s*\**\s*[:\-]\s*(.+?)\s*$"),
    ("company", r"^\**\s*company\s*\**\s*[:\-]\s*(.+?)\s*$"),
]

_CR_BLURB_CACHE: dict[str, str] = {}


def _company_costa_rica_blurb(company: str, location_hint: str = "", *, client: Any = None) -> str:
    """Return a short Haiku-generated paragraph about the company's
    operations in Costa Rica. Cached per-company within the process so a
    16-job run only pays once per unique employer. Falls back to a
    grounded placeholder on any failure (no fabricated facts)."""
    company = (company or "").strip()
    if not company:
        return ""
    if company in _CR_BLURB_CACHE:
        return _CR_BLURB_CACHE[company]

    fallback = (
        f"Public information on {company}'s Costa Rica operations "
        "is not summarized here. Verify the company's local "
        "office, headcount, and hiring practices on its careers "
        "page and on LinkedIn before applying."
    )

    if client is None:
        try:
            from src.llm.client import get_default_client

            client = get_default_client()
        except Exception:
            _CR_BLURB_CACHE[company] = fallback
            return fallback
    _llm = client

    prompt = (
        f"Write a single short paragraph (3–5 sentences, max 110 words) "
        f'describing what is publicly known about "{company}"\'s operations '
        f"in Costa Rica. Cover office presence (San José / Heredia / "
        f"Cartago / GAM / free-trade-zone parks), known business function "
        f"in CR (BPO/shared services/engineering/sales/local-services/etc.), "
        f"approximate years operating in CR, and the general size of the "
        f"local workforce when publicly reported.\n\n"
        f'If "{company}" is a small or local Costa Rican company you\'ve '
        f"only generally heard of, just describe what's publicly known "
        f"(industry, where in CR they operate, rough size) — that's still "
        f"useful. Only if you've genuinely never heard of this company at "
        f"all and have no information about it, respond with the exact "
        f"phrase: NO_VERIFIED_CR_INFO\n\n"
        f"Job posting location hint (use to disambiguate, but do NOT treat "
        f"it as proof of a CR office on its own): {location_hint or '—'}\n\n"
        f"Output rules:\n"
        f"- Plain prose. No bullets, no markdown headings, no code fences.\n"
        f"- Never invent specific numbers, addresses, or named employees.\n"
        f'- Hedge with phrases like "reportedly" or "publicly noted" '
        f"  when you're not certain of a specific fact.\n"
        f"- Maximum 110 words."
    )
    from src.settings import get_settings as _gs

    try:
        content = _llm.complete(
            model=_gs().model_haiku,
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        ).strip()
        if content.startswith("```"):
            content = re.sub(r"^```[a-z]*\n?", "", content)
            content = re.sub(r"\n?```$", "", content)
        if "NO_VERIFIED_CR_INFO" in content.upper():
            _CR_BLURB_CACHE[company] = fallback
            return fallback
        _CR_BLURB_CACHE[company] = content
        return content
    except Exception as e:
        _warn(f"Costa Rica blurb call failed for {company}: {e}")
        _CR_BLURB_CACHE[company] = fallback
        return fallback


def _ensure_markdown_report(
    report: Any, output_dir: Path, job_path: Path, cv_path: Path
) -> Path | None:
    """Build a comprehensive recommendation report. Computes all easy-win
    post-processing in Python, makes ONE bundled Haiku call for the
    LLM-derived strategic insights, then attaches everything to `report`
    under the key "enrichment" before delegating to write_job_report.

    Output path: <output_dir>/<resume_filename>/<job_title>+<company>/report.md

    The resume filename groups all reports generated from the same master CV;
    inside, each job gets its own folder named after the job title + company so
    the user can pair the report with the corresponding cv_proposal.docx.

    Silent-failure guard (§2.2): the target directory is created ONLY when we
    are about to write the markdown report. If any earlier step raises, no
    empty folder is left behind. If a step fails AFTER mkdir, we write a
    _FAILED.md stub so the failure is visible in the output tree.
    """
    target_dir: Path | None = None
    md_path: Path | None = None
    try:
        # Coerce — the agent may have returned a partially-string-typed report.
        report = _as_dict(report) or {}
        job = _as_dict(report.get("job"))
        candidate = _as_dict(report.get("candidate_profile") or report.get("candidate"))

        # Salvage the job + candidate from raw_output when the top-level keys
        # are empty (the agent often wraps everything in JobReport.* envelopes).
        if not job:
            job = _recover_job_from_raw(report)
            if job:
                report["job"] = job
                _info("Recovered job posting from raw_output for the report.")
        if not candidate:
            candidate = _recover_candidate_from_raw(report)
            if candidate:
                report["candidate_profile"] = candidate
                _info("Recovered candidate profile from raw_output for the report.")

        raw_company = (job.get("company") or "").strip()
        title = (job.get("title") or "").strip() or job_path.stem

        # Always stamp the source_file path so the renderer can fall back to
        # parsing it for a "Link: <url>" if the agent didn't populate
        # source_url (per new user request: surface the JOB Link).
        if not job.get("source_file"):
            job["source_file"] = str(job_path)
        if not job.get("source_url"):
            url = _extract_job_link_from_file(job_path)
            if url:
                job["source_url"] = url

        # ── Parse the JD metadata block (Position/Link/Contract/Remote/Company)
        jd_meta = _extract_jd_metadata_from_file(job_path)
        if jd_meta:
            job["jd_metadata"] = jd_meta
            # Backfill canonical fields if the agent missed them.
            if not job.get("title") and jd_meta.get("position"):
                job["title"] = jd_meta["position"]
            if not job.get("company") and jd_meta.get("company"):
                job["company"] = jd_meta["company"]
            if not raw_company and jd_meta.get("company"):
                raw_company = jd_meta["company"].strip()
        report["job"] = job  # write back the enrichment.

        resume_slug = _slugify(cv_path.stem, max_len=60) or "resume"
        job_slug = _slugify(title, max_len=60) or job_path.stem
        company_slug = _slugify(raw_company, max_len=40) if raw_company else ""

        folder_name = f"{job_slug}+{company_slug}" if company_slug else job_slug
        target_dir = output_dir / resume_slug / folder_name
        # NOTE: mkdir deferred to just-before-write below.
        md_path = target_dir / "report.md"

        # ── Pure-Python enrichment (zero LLM cost) ─────────────────────────
        wx = candidate.get("work_experience") or []
        evaluations = report.get("evaluations") or []
        enrichment = {
            "skill_alignment_matrix": _compute_skill_alignment_matrix(job, candidate),
            "quantification_audit": _audit_quantification(wx),
            "weak_words_audit": _audit_weak_words(wx),
            "strong_verbs": _extract_strong_verbs_from_jd(job),
            "pre_submission_checklist": _build_pre_submission_checklist(report),
            "effort_impact_quadrant": _effort_impact_quadrant(report),
            "qualitative_match": _classify_match_with_opus(
                evaluations,
                job,
                _as_dict(report.get("ats_report")),
                _as_dict(report.get("gap_analysis")),
            ),
            "status_badge": _status_badge(report),
            "decision_helper": _decision_helper_with_opus(report),
            "posting_freshness": _posting_freshness(job_path),
            "links": _build_links_section(report, job),
        }

        # ── Bundled Haiku call for the strategic / medium-tier sections ────
        from src.settings import get_settings as _gs

        haiku_model = _gs().model_haiku
        _stage_panel(
            title="🧠 Strategic Insights",
            subtitle="One bundled Haiku call → bullet rewrites, narrative, "
            "application strategy and more",
            items=_STRATEGIC_DELIVERABLES,
            model=haiku_model,
            border="cyan",
        )
        _t0 = time.time()
        with _stage_status("Calling Haiku for strategic insights …", spinner="dots"):
            strategic = _generate_strategic_insights(report)
        _summarize_stage_result(
            title="🧠 Strategic Insights — generated",
            result=strategic,
            items=_STRATEGIC_DELIVERABLES,
            elapsed=time.time() - _t0,
            model=haiku_model,
            border="green" if strategic else "yellow",
        )
        enrichment["strategic"] = strategic  # may be {} on quota / network failure

        # Attach enrichment so downstream helpers (proposal, analysis) can read it.
        report = {**report, "enrichment": enrichment}

        # Stamp TL;DR + effort badge before render — both touch enrichment that
        # the renderer reads. _build_tldr_card depends on decision_helper and
        # ats_report/gap_analysis/effort_estimate; compute effort first.
        enrichment["effort_estimate"] = _estimate_application_effort(report)
        enrichment["tldr"] = _build_tldr_card(report, job_path)

        # Costa Rica company-operations blurb (Haiku, cached per-company).
        company_for_cr = raw_company or _as_dict(job).get("company") or ""
        if company_for_cr:
            enrichment["company_costa_rica"] = _company_costa_rica_blurb(
                company_for_cr,
                location_hint=_as_dict(job).get("location") or "",
            )

        # First mkdir — every artifact from here on must succeed-or-stub.
        target_dir.mkdir(parents=True, exist_ok=True)

        # Copy the source JD file into the output folder as a PDF so the
        # report, proposal, and original job posting all live in one place
        # and share the same .pdf format.
        _copy_job_source_as_pdf(job_path, target_dir)

        # ── Generate the cv_proposal.docx BEFORE writing the markdown report,
        # so the report can include a "Proposal Analysis & Plan" section that
        # describes what the proposal contains and what the candidate should
        # do next.
        proposal_docx = _ensure_cv_proposal(report, target_dir)
        if proposal_docx:
            enrichment["proposal_analysis"] = _build_proposal_analysis(
                report,
                report.get("cv_proposal") or {},
                proposal_docx,
            )

        # ── Write sidecar artifacts (JSON of full report + parsed_resume.json
        # so consumers can sort/filter without re-parsing markdown). §2.6 / §1.8
        _write_sidecar_artifacts(report, target_dir)

        # ── Write the cover letter as a separate file alongside the report.
        # §1.6 — the user asked for a draft, not just an outline.
        _write_cover_letter_sidecar(report, target_dir)

        from src.tools import write_job_report as _wjr

        # Generate the structured Markdown report first — it's the source
        # the PDF renderer will consume. The user's new requirement is to
        # ship a PDF instead of the .md, so we delete the .md after a
        # successful PDF render to keep the output folder clean.
        result = _wjr.func(
            report_json=json.dumps(report, default=str),
            output_path=str(md_path),
        )
        ok = False
        try:
            r = json.loads(result)
            ok = r.get("status") == "ok" and md_path.exists()
        except Exception:
            ok = md_path.exists()
        if not ok or not md_path.exists():
            _write_failed_stub(
                target_dir,
                job_path,
                f"write_job_report returned without writing the .md file. Tool result: {result!r}",
            )
            return None

        # ── Render the PDF from the markdown source ───────────────────────
        pdf_path = target_dir / "report.pdf"
        try:
            from src.pdf_renderer import write_pdf_from_markdown

            md_text = md_path.read_text(encoding="utf-8")
            write_pdf_from_markdown(md_text, pdf_path)
            # Per the user's request: PDF is the user-facing artifact, not
            # the .md. Remove the .md once the PDF is on disk so the
            # output folder stays uncluttered.
            with contextlib.suppress(Exception):
                md_path.unlink()
            report.setdefault("output_files", {})
            report["output_files"]["report_pdf"] = str(pdf_path)
            if _HAS_RICH:
                _console.print(f"[bold green]📄 Report (PDF):[/bold green] [cyan]{pdf_path}[/cyan]")
        except Exception as e:
            import traceback as _tb

            _warn(f"PDF render failed, keeping report.md only: {e}")
            _warn("Traceback:\n" + _tb.format_exc())
            # MD survives — return it as a graceful fallback.
            return md_path
        return pdf_path
    except Exception as e:
        import traceback as _tb

        _warn(f"Could not generate fallback markdown report: {e}")
        tb = _tb.format_exc()
        _warn("Traceback:\n" + tb)
        # If we made it as far as creating the folder, leave a marker so the
        # silent-failure folder doesn't quietly haunt the output tree (§2.2).
        if target_dir is not None and target_dir.exists():
            _write_failed_stub(target_dir, job_path, f"{e}\n\n{tb}")
        return None


def _write_failed_stub(target_dir: Path, job_path: Path, reason: str) -> Path | None:
    """Write a visible failure marker into a job output folder so silent
    failures (empty dirs from earlier crashes) become obvious.

    Writes `_FAILED.pdf` (preferred, matches the new user-facing format)
    with a Markdown fallback only if the PDF renderer is unavailable.
    Returns the path of the artifact actually written, or None."""
    import datetime as _dt

    target_dir.mkdir(parents=True, exist_ok=True)
    ts = _dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    body = (
        f"# ⚠️ Report generation FAILED\n\n"
        f"_Generated: {ts}_\n\n"
        f"**Source job:** `{job_path}`\n\n"
        f"## Reason\n\n```\n{reason.strip()}\n```\n\n"
        f"## What to do\n\n"
        f"- Re-run with `--verbose` to see the full traceback.\n"
        f"- Use `python main.py run --rerun-failed --cv <cv> --jobs <jobs>` "
        f"to retry only the failed/empty folders.\n"
    )
    # Preferred path: PDF, matching the user's "no .md" requirement.
    try:
        from src.pdf_renderer import write_pdf_from_markdown

        pdf_path = target_dir / "_FAILED.pdf"
        write_pdf_from_markdown(body, pdf_path)
        return pdf_path
    except Exception as pdf_e:
        _warn(f"Could not render _FAILED.pdf, falling back to _FAILED.md: {pdf_e}")
    try:
        md_path = target_dir / "_FAILED.md"
        md_path.write_text(body, encoding="utf-8")
        return md_path
    except Exception:
        return None


def _write_sidecar_artifacts(report: dict, target_dir: Path) -> None:
    """Write structured sidecars next to report.md so downstream tooling
    (jq, scripts) can filter/sort without re-parsing markdown. §2.6 / §1.8."""
    try:
        # Full report — sanitized (drop the raw_output dump, which is noisy).
        report_copy = {k: v for k, v in report.items() if k != "raw_output"}
        (target_dir / "report.json").write_text(
            json.dumps(report_copy, indent=2, default=str), encoding="utf-8"
        )
    except Exception as e:
        _warn(f"Could not write report.json sidecar: {e}")
    try:
        cand = _as_dict(report.get("candidate_profile") or report.get("candidate"))
        if cand:
            (target_dir / "parsed_resume.json").write_text(
                json.dumps(cand, indent=2, default=str), encoding="utf-8"
            )
    except Exception as e:
        _warn(f"Could not write parsed_resume.json sidecar: {e}")


def _write_cover_letter_sidecar(report: dict, target_dir: Path) -> Path | None:
    """Render the cover letter body (if available) as cover_letter.pdf next to
    report.pdf so the user gets a ready-to-edit draft. §1.6."""
    cl = _as_dict(report.get("cover_letter"))
    body = (cl.get("body") or cl.get("text") or "").strip()
    if not body:
        return None
    try:
        job = _as_dict(report.get("job"))
        title = (job.get("title") or "Role").strip()
        company = (job.get("company") or "Company").strip()
        opening = (cl.get("opening") or "").strip()
        closing = (cl.get("closing") or "").strip()
        word_count = cl.get("word_count")
        tone = cl.get("tone") or ""
        lines = [f"# Cover Letter — {title} @ {company}\n"]
        meta = []
        if word_count:
            meta.append(f"_~{word_count} words_")
        if tone:
            meta.append(f"_tone: {tone}_")
        if meta:
            lines.append(" · ".join(meta) + "\n")
        if opening and opening not in body:
            lines.append(opening + "\n")
        lines.append(body + "\n")
        if closing and closing not in body:
            lines.append("\n" + closing + "\n")
        md_text = "\n".join(lines)
        path = target_dir / "cover_letter.pdf"
        from src.pdf_renderer import write_pdf_from_markdown

        write_pdf_from_markdown(md_text, path)
        report.setdefault("output_files", {})
        report["output_files"]["cover_letter_pdf"] = str(path)
        return path
    except Exception as e:
        _warn(f"Could not write cover_letter.pdf: {e}")
        return None


def _estimate_application_effort(report: dict) -> dict:
    """Rough effort estimate combining bullets-to-rewrite, ATS keywords
    missing, and gaps to address. Powers the TL;DR card and the cross-job
    leaderboard. §1.10."""
    enrichment = _as_dict(report.get("enrichment"))
    strategic = _as_dict(enrichment.get("strategic"))
    ats = _as_dict(report.get("ats_report"))
    gap = _as_dict(report.get("gap_analysis"))

    bullets_to_rewrite = len(_as_list(strategic.get("bullet_rewrites")))
    keywords_missing = len(_as_list(ats.get("missing_keywords")))
    critical_gaps = len([g for g in _as_list(gap.get("critical_gaps")) if isinstance(g, str)])

    # Rough estimate: 6 min/bullet, 1 min/keyword, 12 min/gap addressed.
    minutes = bullets_to_rewrite * 6 + min(keywords_missing, 12) * 1 + critical_gaps * 12
    minutes = max(minutes, 15)  # never claim less than 15 minutes

    if minutes <= 30:
        band = ("🟢", "quick", f"~{minutes} min")
    elif minutes <= 60:
        band = ("🟡", "moderate", f"~{minutes} min")
    elif minutes <= 120:
        band = ("🟠", "heavy", f"~{minutes} min")
    else:
        band = ("🔴", "very heavy", f"~{minutes} min")

    return {
        "minutes": minutes,
        "band": band[1],
        "emoji": band[0],
        "label": band[2],
        "components": {
            "bullets_to_rewrite": bullets_to_rewrite,
            "keywords_to_inject": keywords_missing,
            "gaps_to_address": critical_gaps,
        },
    }


def _build_tldr_card(report: dict, job_path: Path) -> dict:
    """The 5-line decision card shown at the top of every report. §1.1.
    Derived entirely from data already in `enrichment` — no LLM call."""
    enrichment = _as_dict(report.get("enrichment"))
    decision = _as_dict(enrichment.get("decision_helper"))
    badge = enrichment.get("status_badge") or ("⚪", "Unassessed", "")
    freshness = enrichment.get("posting_freshness") or "—"
    effort = _as_dict(enrichment.get("effort_estimate"))
    gap = _as_dict(report.get("gap_analysis"))
    cf = _as_dict(report.get("consolidated_feedback"))

    crit_gaps = [g for g in _as_list(gap.get("critical_gaps")) if isinstance(g, str)]
    top_blocker = crit_gaps[0] if crit_gaps else None
    if not top_blocker:
        for ev in _as_list_of_dicts(report.get("evaluations")):
            flags = [f for f in _as_list(ev.get("red_flags")) if isinstance(f, str)]
            if flags:
                top_blocker = flags[0]
                break

    top_lever = None
    high_changes = [
        c
        for c in _as_list_of_dicts(cf.get("prioritized_changes"))
        if (c.get("impact") or "").lower() == "high"
    ]
    if high_changes:
        top_lever = high_changes[0].get("description") or None
    if not top_lever:
        for ev in _as_list_of_dicts(report.get("evaluations")):
            strengths = [s for s in _as_list(ev.get("strengths")) if isinstance(s, str)]
            if strengths:
                top_lever = strengths[0]
                break

    return {
        "verdict": decision.get("verdict") or "—",
        "verdict_reason": (decision.get("reason") or "").strip(),
        "badge_emoji": badge[0] if len(badge) > 0 else "⚪",
        "badge_label": badge[1] if len(badge) > 1 else "Unassessed",
        "freshness": freshness,
        "effort_emoji": effort.get("emoji", "⚪"),
        "effort_band": effort.get("band", "—"),
        "effort_label": effort.get("label", "—"),
        "top_blocker": top_blocker,
        "top_lever": top_lever,
    }


def _extract_balanced_object(text: str, start: int) -> str | None:
    """Return the substring containing the first balanced { … } object that
    begins at-or-after `start`. If the input is truncated mid-object, close
    the dangling braces so json.loads can still succeed."""
    open_idx = text.find("{", start)
    if open_idx < 0:
        return None
    depth = 0
    in_string = False
    escape = False
    last_complete_value_end = -1
    for i in range(open_idx, len(text)):
        c = text[i]
        if escape:
            escape = False
            continue
        if c == "\\" and in_string:
            escape = True
            continue
        if c == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return text[open_idx : i + 1]
        elif c == "," and depth >= 1:
            last_complete_value_end = i
    # Truncated. Trim back to the last complete value, then close braces.
    if last_complete_value_end > 0 and depth > 0:
        return text[open_idx:last_complete_value_end] + "}" * depth
    if depth > 0:
        # Naive close — may still fail to parse, but worth one attempt.
        return text[open_idx:] + "}" * depth
    return None


def _parse_raw_output(report: dict) -> dict:
    """Parse the agent's raw_output string (often wrapped in ```json fences
    and sometimes truncated) into a dict. Returns {} on failure.

    Falls back to a brace-balanced recovery so truncated JSON still yields a
    usable dict (with the cut-off branch missing rather than the whole tree)."""
    raw = report.get("raw_output")
    if not isinstance(raw, str) or "{" not in raw:
        return {}
    stripped = re.sub(r"```(?:json)?\s*", "", raw).strip().rstrip("`").strip()
    brace = stripped.find("{")
    if brace > 0:
        stripped = stripped[brace:]
    # Try the cheap parses first.
    for end in (
        len(stripped),
        stripped.rfind("}") + 1,
        stripped.rfind("}", 0, len(stripped) - 1) + 1,
    ):
        if end <= 0:
            continue
        try:
            parsed = json.loads(stripped[:end])
            if isinstance(parsed, dict):
                return parsed
        except Exception:
            continue
    # Truncated JSON — close dangling braces and try again.
    balanced = _extract_balanced_object(stripped, 0)
    if balanced:
        try:
            parsed = json.loads(balanced)
            if isinstance(parsed, dict):
                return parsed
        except Exception:
            pass
    return {}


def _pick_nested(parsed: dict, keys: list[str]) -> dict:
    """Walk a list of dotted-or-flat key paths and return the first non-empty
    dict found. Supports nesting under `JobReport` (the agent's common wrapper)."""
    if not isinstance(parsed, dict):
        return {}
    job_report = parsed.get("JobReport") if isinstance(parsed.get("JobReport"), dict) else {}
    for key in keys:
        for source in (parsed, job_report):
            val = source.get(key)
            if isinstance(val, dict) and val:
                return val
    return {}


def _extract_sub_object(report: dict, key_names: list[str]) -> dict:
    """When the parsed raw_output is unusable (truncated past the key we
    want), fall back to a regex + balanced-brace scan that pulls just one
    sub-object out of the source string."""
    raw = report.get("raw_output")
    if not isinstance(raw, str):
        return {}
    for key in key_names:
        # Match "Key": { … } or 'Key': { … }
        pattern = re.compile(rf'["\']{re.escape(key)}["\']\s*:\s*', re.IGNORECASE)
        m = pattern.search(raw)
        if not m:
            continue
        block = _extract_balanced_object(raw, m.end())
        if not block:
            continue
        try:
            parsed = json.loads(block)
        except Exception:
            continue
        if isinstance(parsed, dict) and parsed:
            return parsed
    return {}


_CANDIDATE_KEYS = ["candidate_profile", "CandidateProfile", "candidate", "Candidate"]
_JOB_KEYS = ["job", "JobPosting", "job_posting", "JobPost"]


def _recover_candidate_from_raw(report: dict) -> dict:
    """Sometimes the final agent emits its JobReport with the candidate
    nested under JobReport.CandidateProfile (or similar) instead of at the
    top level, leaving `candidate_profile` empty after augmentation. As a
    salvage step, try to parse `raw_output` and lift the candidate dict out.

    Returns {} if nothing usable is found.
    """
    parsed = _parse_raw_output(report)
    found = _pick_nested(parsed, _CANDIDATE_KEYS)
    if found:
        return found
    return _extract_sub_object(report, _CANDIDATE_KEYS)


def _recover_job_from_raw(report: dict) -> dict:
    """Salvage the JobPosting dict from raw_output when the top-level `job`
    key is missing or empty (e.g. when the agent wrapped everything inside
    a `JobReport.JobPosting` envelope)."""
    parsed = _parse_raw_output(report)
    found = _pick_nested(parsed, _JOB_KEYS)
    if found:
        return found
    return _extract_sub_object(report, _JOB_KEYS)


def _build_resume_recommendations_text(report: dict) -> dict[str, list[str]]:
    """Re-derive the Must Do / Don't Do recommendations the report renders,
    so we can hand them to Opus as the primary input for the CV proposal.

    This mirrors the structure of `_build_dos_donts` in src/tools.py — kept
    in sync but small enough to duplicate here rather than couple to that
    module's local closure.
    """
    cf = _as_dict(report.get("consolidated_feedback"))
    ats = _as_dict(report.get("ats_report"))
    gap = _as_dict(report.get("gap_analysis"))
    enrichment = _as_dict(report.get("enrichment"))
    weak_audit = enrichment.get("weak_words_audit") or []
    strong = enrichment.get("strong_verbs") or []

    dos: list[str] = []
    donts: list[str] = []

    changes = _as_list_of_dicts(cf.get("prioritized_changes"))
    high = [c for c in changes if (c.get("impact") or "").lower() == "high"]
    for c in high[:5]:
        desc = (c.get("description") or "").strip()
        if desc:
            dos.append(desc)
    if strong:
        dos.append("Mirror these JD verbs: " + ", ".join(_as_list(strong)[:8]))
    missing_kw = _as_list(ats.get("missing_keywords"))
    if missing_kw:
        dos.append("Weave missing ATS keywords naturally: " + ", ".join(missing_kw[:10]))
    for g in _as_list(gap.get("critical_gaps"))[:3]:
        if isinstance(g, str) and g.strip():
            dos.append(f"Address critical gap: {g.strip()}")

    for w in _as_list(weak_audit)[:6]:
        if not isinstance(w, dict):
            continue
        phrases = ", ".join(_as_list(w.get("weak_phrases"))[:3])
        if phrases:
            donts.append(f"Avoid weak phrasing: {phrases}")
    for kw in _as_list(ats.get("overused_keywords"))[:6]:
        donts.append(f"Don't over-repeat: {kw}")
    for ev in _as_list_of_dicts(report.get("evaluations")):
        for rf in _as_list(ev.get("red_flags"))[:2]:
            if isinstance(rf, str):
                donts.append(f"Reviewer flag — {rf}")
    donts.append("Don't claim metrics or scope you can't defend in interview.")

    def _dedup(items: list[str], limit: int) -> list[str]:
        seen, out = set(), []
        for s in items:
            k = s.strip().lower()
            if k and k not in seen:
                seen.add(k)
                out.append(s)
        return out[:limit]

    return {"must_do": _dedup(dos, 10), "dont_do": _dedup(donts, 10)}


def _generate_cv_proposal_with_opus(report: dict, *, client: Any = None) -> dict:
    """Call Opus 4.7 to synthesize an ATS-optimized CV proposal that
    addresses every reviewer's feedback. Returns {} on any failure so the
    caller can degrade gracefully.

    The prompt is now anchored on two PRIMARY inputs — the report's resume
    recommendations (Must Do / Don't Do) and the pre-submission checklist —
    plus supporting context (gap analysis, ATS keywords, reviewer flags).
    """
    if client is None:
        try:
            from src.llm.client import get_default_client

            client = get_default_client()
        except Exception as _e:
            _warn(f"LLM client unavailable — skipping Opus CV proposal: {_e}")
            return {}
    _llm = client

    job = _as_dict(report.get("job"))
    candidate = _as_dict(report.get("candidate_profile") or report.get("candidate"))
    if not candidate:
        # Salvage candidate from the raw agent output (it may have been emitted
        # under a nested JobReport.CandidateProfile rather than top-level).
        candidate = _recover_candidate_from_raw(report)
        if candidate:
            _info("Recovered candidate profile from raw_output for CV proposal.")
            report["candidate_profile"] = candidate

    consolidated = _as_dict(report.get("consolidated_feedback"))
    gap = _as_dict(report.get("gap_analysis"))
    ats = _as_dict(report.get("ats_report"))
    enrichment = _as_dict(report.get("enrichment"))
    strategic = _as_dict(enrichment.get("strategic"))
    evaluations = _as_list_of_dicts(report.get("evaluations"))
    existing_adapted = _as_dict(report.get("adapted_cv"))

    resume_recommendations = _build_resume_recommendations_text(report)
    pre_submission_checklist = list(enrichment.get("pre_submission_checklist") or [])

    feedback_bundle = {
        "resume_recommendations_must_do": resume_recommendations["must_do"],
        "resume_recommendations_dont_do": resume_recommendations["dont_do"],
        "pre_submission_checklist": pre_submission_checklist,
        "executive_summary": consolidated.get("executive_summary", ""),
        "prioritized_changes": [
            {
                "section": c.get("target_section", ""),
                "change_type": c.get("change_type", ""),
                "description": c.get("description", ""),
                "impact": c.get("impact", ""),
            }
            for c in _as_list(consolidated.get("prioritized_changes"))[:15]
            if isinstance(c, dict)
        ],
        "critical_gaps": _as_list(gap.get("critical_gaps"))[:10],
        "minor_gaps": _as_list(gap.get("minor_gaps"))[:8],
        "framing_opportunities": _as_list(gap.get("framing_opportunities"))[:8],
        "missing_ats_keywords": _as_list(ats.get("missing_keywords"))[:30],
        "overused_ats_keywords": _as_list(ats.get("overused_keywords"))[:10],
        "format_issues": _as_list(ats.get("format_issues"))[:10],
        "reviewer_red_flags": [
            {"reviewer": ev.get("agent_role", ""), "flag": rf}
            for ev in evaluations
            for rf in _as_list(ev.get("red_flags"))[:5]
        ][:20],
        "reviewer_recommendations": [
            {"reviewer": ev.get("agent_role", ""), "rec": rec}
            for ev in evaluations
            for rec in _as_list(ev.get("recommendations"))[:5]
        ][:20],
        "strong_verbs_from_jd": _as_list(enrichment.get("strong_verbs"))[:20],
        "career_narrative_hint": strategic.get("career_narrative", ""),
        "bullet_rewrite_examples": _as_list(strategic.get("bullet_rewrites"))[:5],
    }

    cv_snapshot = {
        "full_name": candidate.get("full_name", ""),
        "headline": candidate.get("headline", ""),
        "summary": candidate.get("summary", ""),
        "contact": candidate.get("contact", {}),
        "skills": _as_list(candidate.get("skills"))[:40],
        "languages": _as_list(candidate.get("languages")),
        "certifications": _as_list(candidate.get("certifications")),
        "education": _as_list(candidate.get("education")),
        "work_experience": [
            {
                "company": w.get("company", ""),
                "title": w.get("title", ""),
                "start_date": w.get("start_date", ""),
                "end_date": w.get("end_date", ""),
                "location": w.get("location", ""),
                "bullets": _as_list(w.get("bullets"))[:8],
                "technologies": _as_list(w.get("technologies"))[:15],
            }
            for w in _as_list_of_dicts(candidate.get("work_experience"))
        ],
    }

    target_job = {
        "title": job.get("title", ""),
        "company": job.get("company", ""),
        "location": job.get("location", ""),
        "seniority": job.get("seniority", ""),
        "requirements": [
            r.get("text") if isinstance(r, dict) else r
            for r in _as_list(job.get("requirements"))[:15]
        ],
        "responsibilities": _as_list(job.get("responsibilities"))[:10],
        "ats_keywords": _as_list(job.get("ats_keywords"))[:30],
    }

    prompt = f"""You are an elite resume strategist. Produce an ATS-optimized CV
proposal tailored to the target role. The PRIMARY drivers are the report's
resume recommendations (Must Do / Don't Do) and the pre-submission checklist —
EVERY item below must be reflected in the proposal you return.

Return ONLY a single JSON object matching the schema at the end — no preamble,
no markdown fences, no commentary.

# Hard rules
- DO NOT fabricate employers, dates, titles, education, certifications, or metrics.
- DO NOT invent skills the candidate has not demonstrated. You MAY surface
  latent skills already evidenced by their bullets and reframe them with the
  job's exact terminology.
- If the candidate snapshot includes work_experience or education entries,
  include EVERY entry — do not drop, merge, summarize, or skip any role.
  Reframing bullets, reordering bullets within a role, and adjusting the title
  for clarity are allowed; removing a role is NOT.
- Preserve the original chronological order (most recent first) for both
  experience and education.
- Weave the missing ATS keywords listed in the checklist naturally — never
  keyword-stuff.
- Mirror the JD's strong verbs at the start of bullets when honest to do so.
- Each bullet must be one line, start with a strong verb, and be as quantified
  as the source evidence allows. Do NOT invent percentages or dollar amounts.
- Tone matches the JD's culture (formal vs. casual). Stay first-person-implied
  (no "I"). Use US English.
- Keep summary tight: 3-4 sentences, positioned for THIS role.
- ATS-safe format: single column, no tables, no images, no icons, no graphics,
  no headers/footers. Standard sections only: Summary, Skills, Experience,
  Education, Certifications (only if present), Languages (only if present).
- The `notes_for_candidate` field MUST cite each pre-submission-checklist item
  the proposal already addresses, plus any item the candidate still needs to
  resolve manually (e.g., portfolio link, missing dates).

# PRIMARY INPUT — Resume Recommendations (apply ALL of these)
## Must Do
{json.dumps(resume_recommendations["must_do"], indent=2, ensure_ascii=False)}

## Don't Do
{json.dumps(resume_recommendations["dont_do"], indent=2, ensure_ascii=False)}

# PRIMARY INPUT — Pre-Submission Checklist (resolve every item the proposal can)
{json.dumps(pre_submission_checklist, indent=2, ensure_ascii=False)}

# Target job
{json.dumps(target_job, indent=2, ensure_ascii=False)}

# Candidate's current CV (source of truth — do not exceed its facts)
{json.dumps(cv_snapshot, indent=2, ensure_ascii=False)}

# Supporting agent feedback (gaps, ATS, reviewer flags, JD verbs)
{json.dumps(feedback_bundle, indent=2, ensure_ascii=False)}

# Existing adapted CV draft from earlier agents (may be empty — improve on it)
{json.dumps(existing_adapted, indent=2, ensure_ascii=False) if existing_adapted else "{}"}

# JSON schema to return (exact keys)
{{
  "full_name": "string",
  "contact": {{"email": "string", "phone": "string", "location": "string", "linkedin": "string", "website": "string"}},
  "headline": "Target-role-aligned headline, e.g. 'Brand & Growth Marketing Manager | DTC | LATAM'",
  "summary": "3-4 sentences, positioned for the target role, weaving 4-6 ATS keywords naturally.",
  "skills": ["12-20 deduped skills ordered by relevance to the JD; include any missing keyword the candidate can honestly claim"],
  "experience": [
    {{
      "company": "string",
      "title": "string (may be reframed for clarity but must remain factually true)",
      "location": "string",
      "start_date": "string",
      "end_date": "string",
      "bullets": ["4-6 ATS-friendly bullets starting with strong verbs, mirroring JD language"]
    }}
  ],
  "education": [
    {{"institution": "string", "degree": "string", "field": "string", "start_date": "string", "end_date": "string"}}
  ],
  "certifications": ["only those the candidate already has"],
  "languages": ["e.g. 'English — Fluent', 'Spanish — Native'"],
  "keywords_injected": ["the missing ATS terms you actually wove into the doc"],
  "notes_for_candidate": ["short list of edits the candidate must verify before sending — anything ambiguous, dates to confirm, etc."]
}}"""

    from src.settings import get_settings as _gs

    model_id = _gs().model_opus
    try:
        content = _llm.complete(
            model=model_id,
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}],
        ).strip()
        if content.startswith("```"):
            content = re.sub(r"^```[a-z]*\n?", "", content)
            content = re.sub(r"\n?```$", "", content)
        # Some models still emit a preamble — grab the first {...} block.
        brace = content.find("{")
        if brace > 0:
            content = content[brace:]
        proposal = json.loads(content)
        if not isinstance(proposal, dict):
            return {}
        _backfill_proposal_from_source(proposal, cv_snapshot)
        return proposal
    except Exception as e:
        _warn(f"Opus CV-proposal call failed ({model_id}): {e}")
        return {}


def _backfill_proposal_from_source(proposal: dict, source: dict) -> None:
    """Guarantee education and work experience survive — even if Opus drops or
    truncates them. Restores missing roles by company+title match (so reframed
    titles are honored), but preserves Opus's reordering and bullet rewrites.
    Mutates `proposal` in place.
    """
    # ── Experience: keep what Opus returned, append any source role that's
    # missing (matched case-insensitively by company name, then title).
    src_exp = source.get("work_experience") or []
    out_exp = proposal.get("experience") if isinstance(proposal.get("experience"), list) else []

    def _key(role: dict) -> str:
        c = (role.get("company") or "").strip().lower()
        t = (role.get("title") or "").strip().lower()
        return f"{c}|{t}" if c else t

    out_keys = {
        (role.get("company") or "").strip().lower() for role in out_exp if isinstance(role, dict)
    }
    for src in src_exp:
        if not isinstance(src, dict):
            continue
        company_key = (src.get("company") or "").strip().lower()
        if company_key and company_key in out_keys:
            continue
        # Missing from Opus output — restore verbatim from source.
        out_exp.append(
            {
                "company": src.get("company", ""),
                "title": src.get("title", ""),
                "location": src.get("location", ""),
                "start_date": src.get("start_date", ""),
                "end_date": src.get("end_date", ""),
                "bullets": list(src.get("bullets") or []),
            }
        )
    proposal["experience"] = out_exp

    # ── Education: same strategy, matched by institution.
    src_edu = source.get("education") or []
    out_edu = proposal.get("education") if isinstance(proposal.get("education"), list) else []
    out_inst = {
        (e.get("institution") or "").strip().lower() for e in out_edu if isinstance(e, dict)
    }
    for src in src_edu:
        if not isinstance(src, dict):
            continue
        inst = (src.get("institution") or "").strip().lower()
        if inst and inst in out_inst:
            continue
        out_edu.append(
            {
                "institution": src.get("institution", ""),
                "degree": src.get("degree", ""),
                "field": src.get("field", ""),
                "start_date": src.get("start_date", ""),
                "end_date": src.get("end_date", ""),
            }
        )
    proposal["education"] = out_edu


def _write_cv_proposal_docx(proposal: dict, output_path: Path) -> Path | None:
    """Render an ATS-friendly single-column .docx from the Opus proposal.

    ATS-safe choices:
      - No tables, text boxes, images, or columns.
      - Calibri 11pt body, 14pt name heading, 11pt section headings (bold).
      - Section headers in ALL CAPS — universally parsed by ATS systems.
      - Bullets via the built-in "List Bullet" style.
      - Contact line is a single paragraph, pipe-separated.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        _warn("python-docx not installed — cannot write cv_proposal.docx")
        return None

    if not proposal:
        return None

    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Name (top of CV, centered, bold, 14pt — ATS-safe heading style).
        name = (proposal.get("full_name") or "").strip()
        if name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(16)

        # Contact line.
        contact = proposal.get("contact") or {}
        contact_bits = [
            contact.get(k, "") for k in ("email", "phone", "location", "linkedin", "website")
        ]
        contact_line = " | ".join(b for b in contact_bits if b)
        if contact_line:
            p = doc.add_paragraph(contact_line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        headline = (proposal.get("headline") or "").strip()
        if headline:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(headline)
            run.italic = True

        def _section(title: str) -> None:
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(12)

        summary = (proposal.get("summary") or "").strip()
        if summary:
            _section("Summary")
            doc.add_paragraph(summary)

        skills = [s for s in (proposal.get("skills") or []) if s]
        if skills:
            _section("Skills")
            doc.add_paragraph(", ".join(skills))

        experience = proposal.get("experience") or []
        if experience:
            _section("Experience")
            for role in experience:
                if not isinstance(role, dict):
                    continue
                title = (role.get("title") or "").strip()
                company = (role.get("company") or "").strip()
                location = (role.get("location") or "").strip()
                start = (role.get("start_date") or "").strip()
                end = (role.get("end_date") or "").strip()

                header_left = " — ".join(b for b in [title, company] if b)
                header_right = ""
                if start or end:
                    header_right = f"{start} – {end}".strip(" –")
                if location:
                    header_right = (header_right + f" | {location}").strip(" |")

                p = doc.add_paragraph()
                if header_left:
                    run = p.add_run(header_left)
                    run.bold = True
                if header_right:
                    p.add_run(f"\t{header_right}")

                for bullet in role.get("bullets") or []:
                    if bullet:
                        doc.add_paragraph(str(bullet), style="List Bullet")

        education = proposal.get("education") or []
        if education:
            _section("Education")
            for edu in education:
                if not isinstance(edu, dict):
                    continue
                degree = (edu.get("degree") or "").strip()
                field = (edu.get("field") or "").strip()
                inst = (edu.get("institution") or "").strip()
                start = (edu.get("start_date") or "").strip()
                end = (edu.get("end_date") or "").strip()
                left = ", ".join(
                    b for b in [degree + (f" in {field}" if field else ""), inst] if b.strip()
                )
                date = f"{start} – {end}".strip(" –") if (start or end) else ""
                p = doc.add_paragraph()
                if left:
                    p.add_run(left)
                if date:
                    p.add_run(f"\t{date}")

        certs = [c for c in (proposal.get("certifications") or []) if c]
        if certs:
            _section("Certifications")
            for c in certs:
                doc.add_paragraph(str(c), style="List Bullet")

        languages = [lang for lang in (proposal.get("languages") or []) if lang]
        if languages:
            _section("Languages")
            doc.add_paragraph(", ".join(str(lang) for lang in languages))

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path if output_path.exists() else None
    except Exception as e:
        _warn(f"Could not write cv_proposal.docx: {e}")
        return None


def _ensure_cv_proposal(report: dict, target_dir: Path) -> Path | None:
    """Generate the Opus-driven ATS CV proposal next to report.md."""
    if not target_dir:
        return None
    opus_model = os.getenv("MODEL_OPUS", "claude-opus-4-7")
    _stage_panel(
        title="🧾 ATS CV Proposal",
        subtitle="Synthesizing a tailored, ATS-safe .docx from reviewer "
        "feedback, gap analysis and JD keywords",
        items=_PROPOSAL_DELIVERABLES,
        model=opus_model,
        border="magenta",
    )
    _t0 = time.time()
    with _stage_status("Calling Opus to synthesize the ATS CV proposal …", spinner="dots12"):
        proposal = _generate_cv_proposal_with_opus(report)
    _summarize_stage_result(
        title="🧾 ATS CV Proposal — synthesized",
        result=proposal,
        items=_PROPOSAL_DELIVERABLES,
        elapsed=time.time() - _t0,
        model=opus_model,
        border="green" if proposal else "yellow",
    )
    if not proposal:
        if _HAS_RICH:
            _console.print(
                "[yellow]⚠ No proposal returned — skipping cv_proposal.docx render.[/yellow]"
            )
        return None
    docx_path = target_dir / "cv_proposal.docx"
    with _stage_status("Rendering ATS-safe .docx …", spinner="bouncingBar"):
        written = _write_cv_proposal_docx(proposal, docx_path)
    if written and _HAS_RICH:
        from rich import box as _box
        from rich.panel import Panel as _Panel

        size_kb = written.stat().st_size / 1024 if written.exists() else 0
        body = (
            f"[bold green]✓ ATS CV proposal saved[/bold green]\n"
            f"[dim]📄 Path:[/dim]  [cyan]{written}[/cyan]\n"
            f"[dim]📦 Size:[/dim]  {size_kb:,.1f} KB\n"
            f"[dim]💡 Tip:[/dim]   Open `cv_proposal.docx` next to "
            f"`report.md` to compare against the original CV."
        )
        _console.print(_Panel(body, border_style="green", box=_box.ROUNDED, padding=(1, 2)))
    elif written:
        print(f"   ✓ ATS CV proposal saved: {written}")
    # Stash the structured proposal in the report so the cache preserves it.
    report.setdefault("output_files", {})
    if written:
        report["output_files"]["cv_proposal_docx"] = str(written)
    report["cv_proposal"] = proposal
    return written


def _build_proposal_analysis(report: dict, proposal: dict, docx_path: Path | None) -> dict:
    """Inspect the generated CV proposal and produce a structured analysis
    + action plan the report can render under a dedicated section.

    Compares the proposal against the original parsed CV to surface what
    actually changed (headline, summary, skills added, keywords woven in,
    bullets rewritten), then converts every remaining pre-submission
    checklist item the proposal could not auto-resolve (e.g. portfolio
    link, verification of dates) into a prioritized action plan.
    """
    if not isinstance(proposal, dict) or not proposal:
        return {}

    candidate = _as_dict(report.get("candidate_profile") or report.get("candidate"))
    enrichment = _as_dict(report.get("enrichment"))
    ats = _as_dict(report.get("ats_report"))
    gap = _as_dict(report.get("gap_analysis"))

    # ── Highlights — what's in the proposal at a glance ─────────────────
    highlights: list[str] = []
    summary = (proposal.get("summary") or "").strip()
    headline = (proposal.get("headline") or "").strip()
    skills = [s for s in (proposal.get("skills") or []) if isinstance(s, str)]
    experience = [e for e in (proposal.get("experience") or []) if isinstance(e, dict)]
    education = [e for e in (proposal.get("education") or []) if isinstance(e, dict)]
    keywords_injected = [k for k in (proposal.get("keywords_injected") or []) if isinstance(k, str)]
    notes_from_opus = [n for n in (proposal.get("notes_for_candidate") or []) if isinstance(n, str)]

    if headline:
        highlights.append(f"Repositioned headline: _{headline}_")
    if summary:
        highlights.append(
            f"Rewrote professional summary ({len(summary.split())} words, "
            "front-loaded with ATS keywords)."
        )
    if skills:
        highlights.append(f"Curated {len(skills)} skills ordered by relevance to the JD.")
    if experience:
        bullet_count = sum(len(e.get("bullets") or []) for e in experience)
        highlights.append(
            f"Reframed {len(experience)} role(s) with {bullet_count} ATS-friendly bullets."
        )
    if education:
        highlights.append(
            f"Preserved {len(education)} education entry ({'ies' if len(education) > 1 else 'y'})."
        )
    if keywords_injected:
        highlights.append(
            f"Wove {len(keywords_injected)} missing ATS keyword(s) naturally into the doc."
        )

    # ── Diff vs. original CV (best-effort, only on shared fields) ────────
    diff: list[dict] = []
    src_headline = (candidate.get("headline") or "").strip()
    if src_headline and headline and src_headline.lower() != headline.lower():
        diff.append(
            {
                "field": "Headline",
                "before": src_headline,
                "after": headline,
            }
        )
    src_summary = (candidate.get("summary") or "").strip()
    if src_summary and summary and src_summary != summary:
        diff.append(
            {
                "field": "Summary",
                "before": (src_summary[:200] + "…") if len(src_summary) > 200 else src_summary,
                "after": (summary[:200] + "…") if len(summary) > 200 else summary,
            }
        )
    src_skills = {s.lower() for s in (candidate.get("skills") or []) if isinstance(s, str)}
    added_skills = [s for s in skills if s.lower() not in src_skills]
    if added_skills:
        diff.append(
            {
                "field": "Skills added",
                "before": "—",
                "after": ", ".join(added_skills[:12]),
            }
        )

    # ── Keyword coverage: which missing ATS terms are now covered? ──────
    missing = [k for k in (ats.get("missing_keywords") or []) if isinstance(k, str)]
    proposal_corpus = (
        " ".join([summary, headline, " ".join(skills)])
        + " "
        + " ".join(
            (" ".join(e.get("bullets") or []) + " " + (e.get("title") or "")) for e in experience
        )
    ).lower()
    covered = [k for k in missing if k.lower() in proposal_corpus]
    still_missing = [k for k in missing if k.lower() not in proposal_corpus]

    # ── Action plan — what's left for the candidate to handle manually ──
    checklist = list(enrichment.get("pre_submission_checklist") or [])
    crit_gaps = [g for g in (gap.get("critical_gaps") or []) if isinstance(g, str)]
    action_plan: list[dict] = []

    # P1 — items Opus called out explicitly (portfolio, date verification, etc.)
    for n in notes_from_opus[:8]:
        action_plan.append({"priority": "High", "step": n})

    # P2 — checklist items whose substance is NOT already addressed in the
    # proposal corpus (rough heuristic: at least 4 chars of the item appear).
    def _in_proposal(text: str) -> bool:
        t = (text or "").lower()
        if not t:
            return True
        # If a 12-char snippet of the item appears in the proposal, treat as covered.
        snippet = re.sub(r"[^a-z0-9 ]+", " ", t)
        snippet = " ".join(snippet.split())[:60]
        return bool(snippet) and snippet in proposal_corpus

    for item in checklist:
        if not _in_proposal(item):
            action_plan.append({"priority": "Medium", "step": item})
        if len(action_plan) >= 16:
            break
    # P3 — still-missing keywords (Opus couldn't honestly inject them).
    if still_missing:
        action_plan.append(
            {
                "priority": "Medium",
                "step": "Acquire or surface evidence for: " + ", ".join(still_missing[:8]),
            }
        )
    # P4 — every critical gap that remains is worth interview prep regardless.
    for g in crit_gaps[:4]:
        action_plan.append({"priority": "Low", "step": f"Interview-prep talking point: {g}"})

    # Dedup while preserving priority order.
    seen, deduped = set(), []
    for ap in action_plan:
        key = ap.get("step", "").strip().lower()
        if key and key not in seen:
            seen.add(key)
            deduped.append(ap)
    action_plan = deduped[:20]

    return {
        "docx_path": str(docx_path) if docx_path else "",
        "highlights": highlights,
        "diff": diff,
        "keywords_covered": covered,
        "keywords_still_missing": still_missing[:15],
        "action_plan": action_plan,
    }


# ─── Subcommand: run ──────────────────────────────────────────────────────
